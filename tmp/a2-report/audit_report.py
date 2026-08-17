import json
import re
import sys
from pathlib import Path

from docx import Document


WORD_RE = re.compile(r"[A-Za-z]+(?:['’][A-Za-z]+)?(?:-[A-Za-z]+)*|\b\d+(?:\.\d+)?\b")


def english_part(text: str, style: str) -> str:
    text = text.strip()
    if not text:
        return ""
    if "\n" in text:
        return text.split("\n", 1)[0].strip()
    if " / " in text and (
        style.startswith("Heading")
        or style in {"Figure Caption", "Title", "Subtitle"}
        or re.search(r"[\u3400-\u9fff]", text)
    ):
        return text.split(" / ", 1)[0].strip()
    return text


def main() -> None:
    path = Path(sys.argv[1])
    document = Document(path)
    records = []
    section = "Front matter"
    in_references = False
    totals = {"all": 0, "before_references": 0, "references": 0}
    section_totals = {}

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name
        english = english_part(text, style)
        words = WORD_RE.findall(english)
        if style.startswith("Heading 1"):
            section = english
            if english == "Academic References":
                in_references = True
        count = len(words)
        totals["all"] += count
        if in_references:
            totals["references"] += count
        else:
            totals["before_references"] += count
        section_totals[section] = section_totals.get(section, 0) + count
        records.append(
            {
                "index": index,
                "style": style,
                "section": section,
                "word_count": count,
                "english": english,
                "full_text": text,
            }
        )

    table_words = 0
    table_records = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                english = english_part(text, "Table Text")
                count = len(WORD_RE.findall(english))
                table_words += count
                cells.append({"text": text, "english": english, "word_count": count})
            rows.append(cells)
        table_records.append({"index": table_index, "rows": rows})

    result = {
        "path": str(path.resolve()),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "section_count": len(document.sections),
        "totals": totals,
        "table_words_unclassified": table_words,
        "section_totals": section_totals,
        "records": records,
        "table_records": table_records,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
