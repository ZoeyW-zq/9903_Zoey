from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\GitHub\9903_Zoey")
WORKING_TEMPLATE = ROOT / "tmp" / "a2-report" / "working-template.docx"
DIAGRAM = ROOT / "tmp" / "a2-report" / "narrative-state-network.png"
OUTPUT = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
LIGHT_BLUE = "DCE6F1"
LIGHT_GREY = "F2F2F2"
LIGHT_YELLOW = "FFF2CC"


def font_path(*names):
    fonts = Path(r"C:\Windows\Fonts")
    for name in names:
        path = fonts / name
        if path.exists():
            return str(path)
    raise FileNotFoundError(names)


def rounded_box(draw, xy, fill, outline, radius=22, width=4):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_multiline(draw, box, text, font, fill, spacing=10):
    x1, y1, x2, y2 = box
    max_chars = max(22, int((x2 - x1) / (font.size * 0.58)))
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap(raw, max_chars) or [""])
    rendered = "\n".join(lines)
    bounds = draw.multiline_textbbox((0, 0), rendered, font=font, spacing=spacing, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        rendered,
        font=font,
        fill=fill,
        spacing=spacing,
        align="center",
    )


def arrow(draw, start, end, colour="#5B6573", width=7):
    draw.line([start, end], fill=colour, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 17, y - 26), (x + 17, y - 26)], fill=colour)


def build_diagram():
    image = Image.new("RGB", (1600, 2200), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(font_path("Aptos-Bold.ttf", "arialbd.ttf"), 56)
    node_font = ImageFont.truetype(font_path("Aptos-Bold.ttf", "arialbd.ttf"), 31)
    small_font = ImageFont.truetype(font_path("Aptos.ttf", "arial.ttf"), 27)
    label_font = ImageFont.truetype(font_path("Aptos-Bold.ttf", "arialbd.ttf"), 25)

    draw.text((800, 70), "THE MEMORY ORGANIZER - NARRATIVE STATE NETWORK", font=title_font,
              fill="#111111", anchor="ma")
    draw.text((800, 132), "Freytag phase is identified on every state; arrows show guaranteed progression.",
              font=small_font, fill="#4B5563", anchor="ma")

    nodes = [
        ("EXPOSITION", "OfficeDialogue\nAssistant greeting; learn role, explore, or start work.\nUsed options change, but all routes reconverge.", "#E8F1FB"),
        ("INCITING INCIDENT", "AwaitCrystalBall -> TransitionToHippocampus\nRead the client report, reveal the crystal ball, and enter the subconscious.", "#EAF6EE"),
        ("RISING ACTION", "Hippocampus -> AwaitMemoryPlacement\nInspect and place 3 surface memories in Focus, Context, or Background.\nThe second placement reveals that painful memories are missing.", "#FFF5D9"),
        ("ESCALATION", "GiantCrisis -> SwallowTransition\nThe third placement disturbs the buried memories.\nThe Giant Clown emerges, grabs, and swallows the player.", "#FDE8E7"),
        ("CLIMAX", "MirrorChamber - four mirrors, any order\nEach mirror: approach -> two-round dialogue -> resolve or fail/retry.\nThe fourth resolution dissolves the Giant and releases all deep memories.", "#F7E8F8"),
        ("FALLING ACTION", "FinalMemoryPlacement\nReturn to the Memory Room and redistribute all 7 memories.\nEvery object can be assigned to Focus, Context, or Background.", "#E9F7F5"),
        ("DENOUEMENT", "BackToOffice -> View Report\nThe system composes a qualitative report from the 7 final placements.\nThe report reflects attention without declaring a correct ending.", "#EEF0F4"),
    ]

    left, right = 250, 1350
    y = 190
    boxes = []
    heights = [185, 165, 210, 190, 230, 205, 205]
    for (phase, text, fill), height in zip(nodes, heights):
        box = (left, y, right, y + height)
        boxes.append(box)
        rounded_box(draw, box, fill, "#263238")
        draw.rounded_rectangle((left + 18, y + 18, left + 240, y + 66), radius=15, fill="#263238")
        draw.text((left + 129, y + 42), phase, font=label_font, fill="#FFFFFF", anchor="mm")
        centered_multiline(draw, (left + 60, y + 64, right - 60, y + height - 20), text, node_font, "#111111")
        y += height + 55

    for first, second in zip(boxes, boxes[1:]):
        arrow(draw, ((first[0] + first[2]) / 2, first[3] + 6),
              ((second[0] + second[2]) / 2, second[1] - 8))

    # Local choice loops clarify agency without implying divergent authored endings.
    draw.line([(left - 5, boxes[0][1] + 120), (85, boxes[0][1] + 120), (85, boxes[0][3] - 15), (left - 5, boxes[0][3] - 15)],
              fill="#2E6E9E", width=6)
    draw.polygon([(left - 5, boxes[0][1] + 120), (left - 30, boxes[0][1] + 106), (left - 30, boxes[0][1] + 134)], fill="#2E6E9E")
    draw.text((155, boxes[0][1] + 24), "Explore / learn\nloops", font=small_font, fill="#2E6E9E", anchor="ma")

    draw.line([(right + 5, boxes[4][1] + 125), (1515, boxes[4][1] + 125), (1515, boxes[4][3] - 20), (right + 5, boxes[4][3] - 20)],
              fill="#8A3E8C", width=6)
    draw.polygon([(right + 5, boxes[4][1] + 125), (right + 30, boxes[4][1] + 111), (right + 30, boxes[4][1] + 139)], fill="#8A3E8C")
    draw.text((1450, boxes[4][1] + 25), "Failure returns\nto same mirror", font=small_font, fill="#8A3E8C", anchor="ma")

    draw.text((800, 2085), "Combinatorial outcome layer: 3 attention zones x 7 memories = 3^7 = 2,187 final configurations",
              font=label_font, fill="#1F4E79", anchor="ms")
    image.save(DIAGRAM, dpi=(220, 220))


def remove_body_content(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")
    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    table._tbl.replace(old_grid, new_grid)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, colour="B7B7B7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), colour)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("DDES9903 A2 Report Draft  |  ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, value, end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    style_values = {
        "Title": (24, True, BLACK, 0, 12),
        "Subtitle": (20, False, BLACK, 0, 12),
        "Heading 1": (22, True, BLACK, 14, 7),
        "Heading 2": (15, True, BLUE, 10, 4),
        "Heading 3": (12, True, BLACK, 8, 3),
    }
    for name, (size, bold, colour, before, after) in style_values.items():
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [s.name for s in styles]:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = GREY
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    if "Table Text" not in [s.name for s in styles]:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = "Aptos"
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.0


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.widow_control = True
    return paragraph


def add_placeholder(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RED
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = BLACK
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GREY)
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.add_run(str(value))
    set_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_reference(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(5)
    return p


def build_document():
    doc = Document(WORKING_TEMPLATE)
    remove_body_content(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph(style="Title")
    title.add_run("DDES9903 - Assessment 2 Report Draft")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("The Memory Organizer")

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Normal Table"
    set_table_geometry(callout, [9020])
    set_table_borders(callout, colour="D6B656", size="8")
    set_cell_shading(callout.cell(0, 0), LIGHT_YELLOW)
    cp = callout.cell(0, 0).paragraphs[0]
    cp.style = doc.styles["Table Text"]
    cr = cp.add_run("DRAFT STATUS: ")
    cr.bold = True
    cp.add_run("Complete the highlighted placeholders, confirm the five-minute playthrough target, verify all asset sources, and replace this note before submission.")
    doc.add_paragraph()

    add_heading(doc, "Summary", 1)
    add_table(
        doc,
        ["Submission field", "Draft value"],
        [
            ("WebGL link", "[TO ADD - playable public URL]"),
            ("GitHub repository", "https://github.com/ZoeyW-zq/9903_Zoey"),
            ("First scene", "scene_WebGL"),
            ("Student name / zID", "[TO ADD]"),
        ],
        [2600, 6420],
    )
    add_body(doc, "The Memory Organizer is a short interactive narrative about how attention changes the meaning and emotional force of memory. The player works with a desktop assistant to enter an office worker's subconscious and organise seven memories without deleting or rewriting them. The experience progresses from an office briefing to a Memory Room, a crisis inside a Giant Clown formed from accumulated shame and anxiety, a set of four conversational confrontations, and a final return to the office. Agency is expressed through onboarding pace, the order and response paths of mirror conversations, and especially the physical redistribution of memories among Focus, Context, and Background. The final report reads those placements as a qualitative attention pattern rather than a score or a single correct ending. This preserves a coherent dramatic arc while allowing the player to express an interpretation of the client's present needs.")

    add_heading(doc, "Current implementation and process status", 2)
    add_body(doc, "The five-stage narrative spine is implemented in scene_WebGL.unity through the final compositional report. A script build completed on 16 August 2026 with zero errors. The current work history also shows staged development of office branching, Hippocampus dialogue, memory assignment, the Giant interior, audio, and the final report. The report system stores seven final placements and selects one of three authored fragments for each memory, so 21 fragments support 2,187 possible final configurations. This is a deliberate production strategy rather than a claim that 2,187 hand-authored stories exist.")
    add_table(
        doc,
        ["Evidence in the current project", "Status and report implication"],
        [
            ("Office, surface-memory, mirror, final placement, and report states", "Implemented and scene-bound; supports the full narrative argument."),
            ("Branching dialogue and four mirror conversations", "Implemented; all four mirror orders and failure/retry paths still require full-flow play testing."),
            ("WebGL end-to-end run and in-world report readability", "Outstanding validation; do not describe these as fully tested until evidence is captured."),
            ("XR headset deployment", "XR architecture exists, but deployment evidence is not present in the supplied project context."),
        ],
        [4300, 4720],
    )
    add_placeholder(doc, "[TO ADD BEFORE SUBMISSION: exact measured playthrough duration. The template recommends a maximum of five minutes.]" )

    add_heading(doc, "Narrative State Network", 1)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.keep_with_next = True
    pic.add_run().add_picture(str(DIAGRAM), width=Cm(12.0))
    caption = doc.add_paragraph("Figure 1. Narrative State Network with the Freytag phase assigned to each major state.", style="Figure Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Figure 1 maps the runtime state machine to a Freytag arc. The network separates guaranteed dramatic progression from local agency. The office choices reconverge before entry into the subconscious. Surface-memory placement triggers the crisis regardless of the provisional arrangement, because the Giant is the consequence of buried painful memories rather than a punishment for a wrong answer. The Mirror Chamber is a hub: the four memories can be approached in any order, and harmful responses loop locally without corrupting later states. The fourth successful conversation produces the dramatic climax. The player then receives their widest expressive choice during the falling action by redistributing all seven memories. The denouement is not one authored ending but a report composed from the confirmed attention pattern.")
    add_table(
        doc,
        ["Freytag function", "Narrative state", "Agency carried by the state"],
        [
            ("Exposition", "OfficeDialogue", "Learn the role, explore repeatedly, or start work."),
            ("Inciting incident", "AwaitCrystalBall / TransitionToHippocampus", "Choose when to finish reading and activate the transition."),
            ("Rising action", "AwaitMemoryPlacement", "Inspect and provisionally place three surface memories."),
            ("Escalation", "GiantCrisis / SwallowTransition", "Embodied loss of spatial control while head perspective remains available."),
            ("Climax", "MirrorChamber", "Choose mirror order and dialogue responses; failure creates a retry loop."),
            ("Falling action", "FinalMemoryPlacement", "Revise seven independent Focus/Context/Background assignments."),
            ("Denouement", "BackToOffice / View Report", "Read the personalised consequence pattern and return to page one."),
        ],
        [1850, 2850, 4320],
    )

    add_heading(doc, "Player Agency", 1)
    add_body(doc, "Agency in The Memory Organizer is layered rather than evenly distributed. At the office, the player controls onboarding pace. They may hear the job explanation, explore, return, and then begin; a used explanation option disappears while exploration remains repeatable. These routes create at least five intentional onboarding sequences, but they share the same client assignment. The design therefore offers freedom to establish curiosity and readiness without multiplying downstream production work.")
    add_body(doc, "The strongest authored branching occurs inside the Giant. Four mirrors can be visited in any order, already producing 24 orderings before dialogue variation is counted. Each mirror has a two-round conversation. Two plausible first responses continue into different second-round prompts; harmful shortcuts at either round end the attempt, intensify the local voice and environment, and restore Start Conversation. A failure changes the player's immediate experience but does not permanently remove content or generate an incoherent global state. This makes mistakes narratively legible as the persistence of an internal belief rather than as a conventional game-over screen.")
    add_body(doc, "The final redistribution is the principal form of expressive agency. Every memory can be moved among Focus, Context, and Background, and the arrangement can be revised until confirmation. Seven independent ternary choices yield 3^7, or 2,187, attention configurations. These are real state differences: the report names each final category and selects a corresponding outcome fragment for every object. However, the system avoids a simplistic good/bad meter. Focus is not always positive, and Background is not equivalent to deletion. The player is choosing what should occupy current attention, not editing the client's history.")
    add_body(doc, "This approach deliberately combines real and illusory choice. The Giant crisis always occurs after the third surface memory, so the preliminary layout does not branch around the central conflict. That convergence is visible only at the system level; within the story, the crisis follows the discovery that painful memories have been excluded. The illusion of an unconstrained story world is supported by free inspection order, spatial revision, local dialogue loops, and a final state that genuinely changes the report. The design is innovative in its use of compositional consequence: a small authored vocabulary produces many coherent interpretations without pretending to simulate an unlimited story.")
    add_body(doc, "A limitation remains. During the grab and swallow sequence, translational movement is intentionally locked to prevent the player escaping the staged transition, although camera rotation and tracked head orientation are not forcibly overridden. This is a short restriction on navigation, not a detached screen-based cutscene. The final report should describe it accurately rather than claim uninterrupted locomotion.")

    add_heading(doc, "Narrative Coherence", 1)
    add_body(doc, "The design begins with a fixed thematic rule: a Memory Organizer can redistribute attention but cannot delete, rewrite, or manufacture memories. This rule protects causality across every route. The client has internalised the belief that worth depends on perfect performance, endurance, and not burdening others. Three recent memories and four deep memories test that belief from different angles. The water bottle repeats working while ill; the sunset and LEGO memories offer alternative sources of value; the medal, clock, phone, and correction pen materialise perfectionism, guilt, isolation, and fear of mistakes. Because every object relates to the same belief system, different inspection orders remain coherent.")
    add_body(doc, "Freytag structure is maintained by separating the dramatic spine from variable content. Exposition and the inciting transition establish the role and rules. Surface-memory placement raises questions about what is missing. The Giant's emergence escalates the conflict, and four mirror conversations make the client confront the beliefs sustaining it. The fourth resolution is the climax because it removes the Giant's support. Final redistribution is falling action: the player applies new understanding rather than fighting a second antagonist. The office report supplies denouement by reflecting the chosen attention pattern.")
    add_body(doc, "Several gates prevent incoherent sequencing. The crystal ball appears only after Start Work dialogue and client-report reading. The initial crisis begins only after three distinct surface objects are placed. A mirror contributes to progression only once, after a constructive second-round response. Final confirmation requires all seven distinct memories. The report reads only the confirmed final placement dictionary. These constraints make state transitions explicit while leaving order and interpretation flexible inside each state.")
    add_body(doc, "The design also limits production burden. Four mirror instances share one data-driven MemoryDialogueController rather than separate scripts. Initial and final placement reuse one phased MemoryPlacementController. Most importantly, the report uses 21 object-level fragments rather than 2,187 complete reports. This compositional structure is robust because every fragment refers to one memory and one attention category; combining them cannot contradict a global success label because no such label exists. It is both an implementation economy and a narrative choice that resists reducing mental health to a score.")

    add_heading(doc, "Use of Space", 1)
    add_body(doc, "Space carries narrative meaning before it serves navigation. The ordinary office frames memory work as a professional task and gives the assistant a diegetic place to brief the player. The crystal ball acts as a threshold object between administrative reality and the client's subconscious. The Memory Room is an external model of attention: Focus, Context, and Background are not menu options but physical zones. Moving an object changes its relationship to the player's body and to other memories, making an abstract cognitive judgment spatially inspectable.")
    add_body(doc, "The story then changes scale. The Giant Clown breaks the apparent safety of the Memory Room, removes the roof, and swallows the player. Inside the Giant, four mirrors spatially separate painful beliefs while enclosing them in one body. The player can approach the mirrors in any order, so the chamber functions as a narrative hub rather than a corridor. When all four are resolved, the space dissolves and the freed objects return to the damaged Memory Room. The damaged ceiling makes the crisis leave a persistent spatial trace even though the plot has moved into falling action.")
    add_body(doc, "Spatial agency is bounded where coherence requires it. Placement zones permit repeated revision, while non-overlapping trigger volumes and distinct object identities ensure that each final memory has one readable state. Coarse Office, Hippocampus, and Nightmare roots are activated by the state machine so only the relevant story space is live, reducing runtime cost for XR. The current build still needs Play Mode checks for overlapping zone triggers, direct movement between zones, return-point alignment, and player visibility at the target WebGL resolution.")

    add_heading(doc, "Use of Sensory Triggers (or Multimodal Elements)", 1)
    add_body(doc, "Audio, lighting, motion, and fades are used as state information rather than decoration. Picking up a memory plays its associated account, joining object handling to voice. After the second surface placement, a low sound, light fluctuation, and room disturbance signal that the visible set is incomplete. During the crisis, footsteps finish before rumble begins, the Giant approaches, and the roof and player are moved toward the same hand anchor. These cues turn an abstract accumulation of pressure into an approaching physical threat.")
    add_body(doc, "Transitions use different visual languages. White fade accompanies the crystal-ball transfer into the subconscious, suggesting connection and entry. Black fade overlaps the final movement toward the Giant's mouth and protects the discontinuity of teleporting to the swallow/fall space. Global Volume weights distinguish Office, Hippocampus, and Nightmare states; the Nightmare treatment fades out as the four mirrors are resolved. Particle emission also fades when the crisis begins, helping the environment communicate a systemic change rather than relying only on subtitles.")
    add_body(doc, "Mirror audio creates persistence. Each unresolved mirror loops a negative internal statement. A harmful response intensifies the voice and nearby environment, then returns the mirror to idle; a constructive resolution stops the echo and releases the original object. The player can therefore hear the chamber become quieter as progress accumulates. In VR, holding the crystal ball can also provide controller haptics; the WebGL route substitutes a click interaction. This platform distinction should be demonstrated with separate evidence rather than described as identical.")
    add_body(doc, "Multimodal cues are sequenced by the same state controllers that govern narrative progression, which reduces the risk of contradictory signals. Remaining validation is material: looping footsteps could stall the crisis, subtitle/audio timing must be checked during the black transition, and the longest report layouts must be tested for readability. These are current risks, not completed results.")

    add_heading(doc, "Use of Semiotic Elements", 1)
    add_body(doc, "The central semiotic system is the relationship between memory objects and attention. The objects are deliberately ordinary: a used water bottle, a sunset photograph, LEGO bricks, a second-place medal, an alarm clock, an old phone, and a red correction pen. Their familiarity lets personal meaning emerge through context. The bottle signifies endurance at the expense of health; the clock turns rest into guilt; the unsent message on the phone signifies fear of burdening others; the red pen turns one mistake into a verdict on competence. The medal and LEGO model contrast external ranking with patient private achievement, while the sunset represents value that does not need productivity or approval.")
    add_body(doc, "The Giant Clown combines the client object's existing clown imagery with anxiety, shame, pressure, and self-criticism accumulated beyond manageable scale. It is not a random monster and not the consequence of a wrong player choice. The mirrors signify repeated internal speech: they return the client's beliefs as apparently objective truths. They shatter automatically after a successful conversation because the task is reinterpretation, not physical aggression. The original object remains, reinforcing the rule that the memory has not been erased.")
    add_body(doc, "Focus, Context, and Background avoid moral labels. Their names describe attentional proximity, and the final report uses qualitative coloured labels without success/failure colours, stars, or percentages. This matters because the same placement can have mixed implications. Putting a painful memory in Background can reduce immediate burden but also make its lesson less available; putting a positive memory in Focus can be restorative, but the system does not declare a universal optimum. The semiotic system therefore supports agency while protecting the story's ethical claim that painful and pleasant experiences continue to coexist.")

    add_heading(doc, "Use of Embodied Actions and/or Enacted Elements", 1)
    add_body(doc, "The player learns the narrative by acting on its metaphors. They inspect the office, activate the crystal ball, pick up memory objects, listen while holding them, place them in spatial attention zones, approach mirrors, begin conversations, and revise the final arrangement before confirmation. These actions make cognition tangible: attention becomes distance and placement; confrontation becomes approaching a reflective surface; reinterpretation releases an object that can later be handled again.")
    add_body(doc, "Embodiment also changes the dramatic stakes. The Giant removes the roof, reaches for the player, carries them toward its mouth, and initiates a swallow/fall transition. The system moves the XR origin by position while preserving tracked head orientation, so the player remains perceptually present rather than watching an external avatar. Movement locking is shared across VR and WebGL: locomotion stops during the grab, but mouse look or head turning remains available. This is an authored loss of agency that supports the story's crisis, followed by a return of agency in the Mirror Chamber and the widest expressive control during final redistribution.")
    add_body(doc, "The assistant also enacts narrative change. It briefs the player, notices missing memories, reacts to the crisis, regains a teleport signal after all four echoes stop, and returns the player and freed objects to the Memory Room. Because these actions are tied to completion events, the assistant behaves as a participant in the story world rather than a floating instruction panel. Diegetic computer pages, world-space choices, mirrors, buttons, and physical zones further reduce reliance on screen-fixed UI.")
    add_placeholder(doc, "[TO ADD: XR headset test evidence, device name, comfort observations, frame-rate result, and one labelled screenshot. Without this evidence, do not claim that c7 is satisfied.]" )

    add_heading(doc, "Collaborative Practice", 1)
    add_placeholder(doc, "[SECTION RESERVED FOR THE STUDENT. Insert English screenshots from official class channels for Weeks 6-10. Caption every item with the exact rubric label (d1-d12), date, channel, what you contributed, and what changed as a result.]" )
    add_table(
        doc,
        ["Rubric label", "Date and channel", "Evidence screenshot", "Why the contribution was substantive / reciprocal"],
        [
            ("[d__]", "[TO ADD]", "[INSERT ENGLISH SCREENSHOT]", "[TO ADD]"),
            ("[d__]", "[TO ADD]", "[INSERT ENGLISH SCREENSHOT]", "[TO ADD]"),
            ("[d__]", "[TO ADD]", "[INSERT ENGLISH SCREENSHOT]", "[TO ADD]"),
        ],
        [1300, 1900, 2500, 3320],
    )

    add_heading(doc, "Academic References", 1)
    add_reference(doc, "Dourish, P. (2001). Where the Action Is: The Foundations of Embodied Interaction. MIT Press.")
    add_reference(doc, "Freytag, G. (1894). Freytag's Technique of the Drama (E. J. MacEwan, Trans.). Scott, Foresman and Company. (Original work published 1863).")
    add_reference(doc, "Jenkins, H. (2004). Game design as narrative architecture. In N. Wardrip-Fruin and P. Harrigan (Eds.), First Person: New Media as Story, Performance, and Game (pp. 118-130). MIT Press.")
    add_reference(doc, "Murray, J. H. (1997). Hamlet on the Holodeck: The Future of Narrative in Cyberspace. MIT Press.")
    add_reference(doc, "Ryan, M.-L. (2001). Narrative as Virtual Reality: Immersion and Interactivity in Literature and Electronic Media. Johns Hopkins University Press.")
    add_reference(doc, "Slater, M. (2009). Place illusion and plausibility can lead to realistic behaviour in immersive virtual environments. Philosophical Transactions of the Royal Society B, 364(1535), 3549-3557. https://doi.org/10.1098/rstb.2009.0138")
    add_placeholder(doc, "[CHECK REQUIRED: align citation style with course requirements and add the specific Week 2-4 and Week 8 course materials actually used.]" )

    add_heading(doc, "Asset Citations", 1)
    add_body(doc, "The following sources are visible in the repository documentation. The final submission should include only assets that remain in the submitted build, plus the exact source and licence for every imported model, image, sound, font, and animation.")
    for source in [
        "EZPZ Interaction Toolkit, Matt Cabanag, MIT licence: https://www.linkedin.com/in/mattavc/",
        "Unity Starter Assets - First Person Character Controller: https://assetstore.unity.com/packages/essentials/starter-assets-first-person-character-controller-196525",
        "RN Low Poly Dungeon Lite: https://assetstore.unity.com/packages/slug/224090",
        "Freesound source listed by the toolkit: https://freesound.org/people/sandyrb/sounds/105348/",
        "Freesound source listed by the toolkit: https://freesound.org/people/kwahmah_02/sounds/275072/",
    ]:
        add_reference(doc, source)
    add_placeholder(doc, "[TO VERIFY AND CITE: Whiteclown N Hallin model/animations; Ace Combat 7 Medals ACE model; clock; cell phone; breakable glass; BrickToy assets; all memory images; all assistant, memory, ambience, heartbeat, rumble, footsteps, monster, and swallow audio files. Remove unused inherited sample citations.]" )

    add_heading(doc, "AI Acknowledgements", 1)
    add_body(doc, "OpenAI Codex was used as a development and writing assistant. For this report draft, Codex read the project context and assessment documents, compared the report structure with the detailed marking sheet, generated an initial Narrative State Network diagram, and drafted prose for later student review. AI output was treated as a draft: project claims were restricted to evidence present in the repository, and unknown links, collaboration evidence, asset provenance, timing, and headset test results were left as placeholders.")
    add_table(
        doc,
        ["Service", "Use", "Prompt / record to retain"],
        [
            ("OpenAI Codex", "Report structure, first draft, diagram, document formatting", "User prompt dated 16 Aug 2026: read PROJECT_CONTEXT and the three assessment files, then draft the report; leave collaboration for later."),
            ("OpenAI Codex", "Unity implementation, debugging, tests, or scene analysis", "[PASTE THE EXACT RELEVANT PROMPTS OR PROVIDE A LINK/APPENDIX RECORD]"),
        ],
        [1800, 2800, 4420],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_diagram()
    build_document()
