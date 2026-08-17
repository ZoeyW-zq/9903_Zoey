from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(r"E:\GitHub\9903_Zoey")
SOURCE = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer_Bilingual_Actual_Outcomes_Stage3_Revised.docx"
OUTPUT = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer_Bilingual_Actual_Outcomes_Corrected.docx"

CN_FONT = "Microsoft YaHei"
CN_GREY = RGBColor(0x58, 0x62, 0x70)


REVISIONS = {
    5: (
        "Memory Organizer is a short immersive interactive story about repairing an anxious Dreamer's relationship with memory and attention. The player begins as a newly hired Memory Organizer in a warm office, enters a hippocampus-inspired memory room, and provisionally places three recent memories in Focus, Context, or Background. This apparently manageable task is interrupted when a giant clown tears open the room, captures the player, and swallows them. The swallow forms the climax of the story. The player then enters a compact body-interior space designed for one person, where four containers hold a second-place medal, old alarm clock, old phone, and red correction pen. This stage is the falling action shown in Figure 1. The player approaches the containers, navigates two-round conversations, and receives different dialogue responses to constructive or harmful choices. A constructive resolution breaks the relevant container and releases the memory; a harmful response leaves the memory unresolved and allows another attempt. After all four objects are released, the player returns to the memory room for the resolution: freely redistributing all seven memories across Focus, Context, and Background. These attention choices produce real psychological consequences for the client, which are assembled into the final report before the neutral Close Session ending.",
        "《记忆整理师》是一段关于修复焦虑梦境者与记忆及注意力关系的短篇沉浸式互动故事。玩家以新入职的记忆整理师身份从温暖的办公室出发，进入以海马体为灵感的记忆空间，并把三段近期记忆暂时放入“焦点”“语境”或“背景”。这个看似可控的任务被巨型小丑打断：它撕开房间、抓住玩家并将其吞下。吞咽构成故事的高潮。随后，玩家进入一个适合单人活动的紧凑身体内部空间，四个容器分别封存第二名奖牌、旧闹钟、旧手机和红色批改笔；这一阶段对应图 1 中的下降行动。玩家接近容器并完成两轮对话，建设性与有害选择会收到不同的话语回应。建设性解决会使对应容器破碎并释放记忆；有害回应则让记忆保持未解决状态，并允许玩家重新尝试。四个物件全部释放后，玩家回到记忆空间进入解决阶段：自由地把七段记忆重新分配到“焦点”“语境”和“背景”。这些注意力选择会对来访者心理产生真实影响，并被组合进最终报告，之后以中性的“Close Session”结束。",
    ),
    9: (
        "Figure 1 defines the dramatic stages used in this report. The office and initial memory placement establish the exposition and build tension. GiantCrisis and SwallowTransition form the climax: the safe room is broken open, movement is restricted, and the player is captured and swallowed. Entering the Giant's body begins the falling action rather than extending the climax. In this compact space, four deep memories are sealed in four containers. The player may choose which container to approach and selects responses in a two-round conversation, but every path remains inside the same overall story direction. A harmful choice produces a different negative reply and leaves that container available for another attempt; it does not create a game-over or alter the final report. A constructive response gives the memory a new interpretation, breaks the container, and releases the object. After all four objects are free, FinalMemoryPlacement provides the resolution by letting the player redistribute all seven memories. The office report then acts as the denouement by showing the psychological consequences of those confirmed attention choices.",
        "图 1 规定了本报告采用的戏剧阶段。办公室与初次记忆放置建立开端并逐步累积张力。GiantCrisis 与 SwallowTransition 构成高潮：安全空间被破坏，玩家的移动受到限制，随后被抓住并吞下。进入巨人体内后开始的是下降行动，而不是高潮的延续。在这个紧凑空间中，四段深层记忆被封存在四个容器里。玩家可以选择先接近哪个容器，并在两轮对话中选择回应，但所有路径仍处于同一条总体剧情方向中。有害选择只会得到不同的负面话语回应，并让该容器保持可重新尝试；它不会触发 Game Over，也不会改变最终报告。建设性回应会为记忆提供新的理解，使容器破碎并释放物件。四个物件全部恢复自由后，FinalMemoryPlacement 通过允许玩家重新分配七段记忆来完成解决阶段。办公室报告随后作为结局，展示已确认注意力选择对来访者心理产生的后果。",
    ),
    12: (
        "Player agency is intentionally limited by the subject and dramatic structure. The story concerns a Memory Organizer confronting the client's anxiety, so the player cannot reject the assignment, prevent the Giant's appearance, avoid being swallowed, or create a completely different plot. The compact body-interior space also offers limited spatial expansion rather than open exploration. Early office choices, provisional placement, container order, and dialogue responses change the rhythm and local experience, but they reconverge on the same required sequence. Acknowledging these limits is important because the project does not claim to simulate an unrestricted branching world.",
        "玩家能动性受到题材与戏剧结构的有意限制。故事讲述记忆整理师面对来访者的焦虑，因此玩家不能拒绝任务、阻止巨人出现、避免被吞下，也不能创造一条完全不同的剧情。紧凑的身体内部空间同样只提供有限的空间拓展，而不是开放式探索。办公室前期选择、临时放置、容器顺序和对话回应会改变节奏与局部体验，但最终仍会汇合到同一条必要流程。承认这些限制十分重要，因为项目并不声称自己模拟了一个没有边界的分支世界。",
    ),
    13: (
        "The four-container stage provides local, not dominant, agency. The player can decide which deep memory to address first and navigate two rounds of dialogue for each container. Round 1 presents three choices: two constructive interpretations lead to different Round 2 prompts, while a harmful shortcut receives a different negative reply and ends that attempt. Each Round 2 branch again distinguishes a constructive response from a harmful one. These choices let the player participate in reframing achievement, rest, help-seeking, and mistakes, but they do not create separate long-term storylines. Every unresolved container can be retried, and all four memories must eventually be released before the story can continue.",
        "四容器阶段提供的是局部能动性，而不是项目中最主要的能动性。玩家可以决定先处理哪段深层记忆，并为每个容器完成两轮对话。第一轮提供三个选项：两个建设性解释分别进入不同的第二轮问题；有害捷径会收到不同的负面回应，并结束本次尝试。每条第二轮分支也再次区分建设性回应与有害回应。这些选择让玩家参与重新理解成就、休息、求助和错误，但不会形成彼此独立的长期剧情线。每个尚未解决的容器都可以重试，而且故事必须在四段记忆全部释放后才能继续。",
    ),
    14: (
        "Failure in the body-interior chamber is communicated through dialogue only. A harmful choice triggers its authored response, closes the current attempt, and leaves the container unresolved so the player can try again. Success is communicated through a physical state change: once the memory receives a constructive interpretation, the container breaks and releases the original object. The contrast between an unchanged unresolved container and a broken resolved container gives the player clear feedback without adding a score or permanent punishment.",
        "身体内部空间中的失败只通过对话表达。有害选择会触发预先写好的回应，结束当前尝试，并让容器保持未解决状态，使玩家可以再次尝试。成功则通过物理状态变化表达：当记忆获得建设性的新理解后，容器会破碎并释放原始物件。保持完整的未解决容器与已经破碎的解决容器形成对比，为玩家提供清晰反馈，同时不加入分数或永久惩罚。",
    ),
    15: (
        "The strongest player agency appears in the final attention allocation. The player can freely move each of the seven memories among Focus, Context, and Background, revise the arrangement before confirmation, and decide how much present attention every object should receive. Seven independent ternary choices create 3^7 = 2,187 possible attention patterns. More importantly, these are not cosmetic variations: the selected position changes the psychological interpretation written for that memory, and the seven selected outcomes are combined in the client's final report. The player therefore cannot rewrite the client's past, but can materially influence how strongly each experience shapes the client's present attention, self-understanding, and recovery.",
        "玩家最强的能动性体现在最终注意力分配。玩家可以自由地把七段记忆移动到“焦点”“语境”和“背景”，在确认前反复修改安排，并决定每个物件在当下应获得多少注意力。七个彼此独立的三元选择产生 3^7，即 2,187 种可能的注意模式。更重要的是，这些并非装饰性差异：选定位置会改变该段记忆对应的心理解释，七条选中结果会共同组成来访者的最终报告。因此，玩家不能改写来访者的过去，但能够实质性影响每段经历在多大程度上塑造来访者当下的注意力、自我理解与恢复过程。",
    ),
    16: (
        "Control changes support this hierarchy of agency. During the Giant's grab, translational movement is locked so the scripted climax cannot be escaped, while head tracking or mouse look remains available. SwallowController then transfers the player through the black-screen fall into the body-interior space, where movement is restored for approaching the four containers. After the four memories are released, the transfer back to the memory room leads to the most open interaction in the project: the final seven-object arrangement. The progression therefore moves from temporary loss of control, through limited local choices, to a final decision system with genuine downstream consequences.",
        "控制权变化支撑了这一能动性层级。巨人抓取期间，位移会被锁定，使编排好的高潮无法被走位逃脱，但仍保留头部追踪或鼠标观察。SwallowController 随后通过黑屏坠落把玩家送入身体内部空间，并恢复移动，使玩家可以接近四个容器。四段记忆释放后，玩家被送回记忆空间，进入项目中最开放的交互：最终七物件安排。因此，体验从暂时失去控制，经过有限的局部选择，最终到达一个具有真实后续影响的决策系统。",
    ),
    22: (
        "Narrative coherence is protected by one ethical and causal rule: a Memory Organizer may redistribute attention but cannot delete, rewrite, or invent memories. The four body-interior containers hold connected deep memories rather than unrelated side stories. The medal carries the belief that only first place proves worth; the clock makes unfinished work turn rest into guilt; the phone connects previous rejection with fear of asking for help; and the red pen turns one public mistake into a lasting judgment of ability. These memories explain the client's current anxiety and why rearranging only the three recent objects cannot resolve the problem. The player must first help each deep memory gain a new interpretation before all seven objects can be reconsidered together.",
        "叙事连贯性由一条伦理与因果规则保护：记忆整理师可以重新分配注意力，但不能删除、改写或制造记忆。身体内部的四个容器封存的是彼此关联的深层记忆，而不是互不相关的支线故事。奖牌承载“只有第一名才能证明价值”的信念；闹钟使未完成工作把休息转化为内疚；手机把曾经被拒绝与害怕求助连接起来；红笔则把一次公开犯错变成对能力的长期判断。这些记忆解释了来访者当前的焦虑，也说明为什么只重新安排三个近期物件无法解决问题。玩家必须先帮助每段深层记忆获得新的理解，之后才能把全部七个物件放在一起重新考虑。",
    ),
    23: (
        "The report follows the same Freytag labels as Figure 1. Exposition introduces the role and rules; tension building uses the first three memory objects; the Giant's appearance, capture, and swallow form the climax because they produce the greatest loss of safety and control. The body-interior container conversations are the falling action: the immediate threat has already peaked, and the player now works through the beliefs that supported it. Each constructive interpretation breaks one container and releases one memory, gradually moving the story away from crisis. Returning to the memory room and allocating attention across all seven objects provides the resolution. The generated office report is the denouement because it shows the consequences of the player's final decisions.",
        "本报告采用与图 1 完全一致的 Freytag 标签。开端介绍角色与规则；张力累积阶段使用前三个记忆物件；巨人的出现、抓取与吞咽构成高潮，因为它们带来最强烈的安全感与控制权丧失。身体内部的容器对话属于下降行动：即时威胁已经到达顶点，玩家开始处理支撑这种威胁的信念。每一次建设性的新理解都会打破一个容器、释放一段记忆，使故事逐步远离危机。回到记忆空间并为七个物件分配注意力构成解决阶段。办公室中生成的报告则是结局，因为它展示玩家最终决定产生的后果。",
    ),
    24: (
        "State gates keep this mostly linear structure coherent. The climax begins only after three unique recent objects are placed. SwallowTransition moves the player into the body-interior scene before movement is restored. Each container contributes to progress only once and only after a constructive Round 2 response; a harmful response gives its alternate dialogue and leaves that container available for another attempt. Completing one container cannot accidentally end the stage, and retrying an unresolved container cannot duplicate progress. When all four are complete, the space breaks down and the released objects return to explicit positions in the memory room with gravity and holdability restored. The final report can be confirmed only after all seven objects receive attention zones.",
        "状态门控使这条以线性为主的结构保持连贯。三个不同的近期物件都完成放置后，高潮才会开始。SwallowTransition 先把玩家送入身体内部场景，之后再恢复移动。每个容器只能贡献一次进度，而且必须在第二轮选择建设性回应；有害回应只会给出另一段话语，并让该容器保持可重新尝试。完成一个容器不会意外结束整个阶段，重试未解决容器也不会重复增加进度。当四个容器全部完成后，空间开始解体，释放出的物件回到记忆空间中明确的位置，并恢复重力与可持有状态。只有全部七个物件都获得注意区域后，最终报告才能确认。",
    ),
    25: (
        "The later stage remains practical because four differently authored container conversations share one data-driven dialogue controller and the same proximity, choice, retry, completion, and break logic. Their differences are configured as dialogue text, audio, Round 2 branches, and constructive or harmful responses rather than four separate code paths. The initial and final placement stages likewise share one phased controller. Most of the project's meaningful variation is concentrated in the report system: seven ternary placements select seven of 21 object-specific psychological outcomes, supporting 2,187 coherent combinations without pretending that the main plot itself contains 2,187 branches.",
        "后半段仍然具有可实现性，因为四场内容不同的容器对话共享一个数据驱动的对话控制器，以及相同的接近、选择、重试、完成与破碎逻辑。它们的差异通过对话文字、音频、第二轮分支和建设性或有害回应进行配置，而不是写成四套独立代码。初始与最终放置阶段也复用同一个分阶段控制器。项目中最有意义的变化主要集中在报告系统：七个三元位置会从 21 条物件专属心理结果中选择七条，从而支持 2,187 种连贯组合，同时不假装主线剧情本身拥有 2,187 条分支。",
    ),
    27: (
        "The project uses distinct spaces to control narrative scale, but spatial exploration remains deliberately limited. The office is a warm client-facing care space, while the hippocampus room externalises attention through a central worktable and three physical placement zones. The crystal ball provides a clear threshold into the Dreamer's subconscious. The Giant then breaks the safety of the memory room and carries the player into a separate body-interior space. This later space is compact and intended for one player rather than designed as a large open exploration environment. Its purpose is to focus attention on the four sealed memories and their conversations, not to create extensive navigation content.",
        "项目使用不同空间来控制叙事尺度，但空间探索受到有意限制。办公室是温暖、面向来访者的关怀空间；海马体记忆空间则通过中央工作桌和三个实体放置区域，把注意力外化。水晶球为进入梦境者潜意识提供清晰门槛。随后，巨人打破记忆空间的安全感，并把玩家带入独立的身体内部空间。这个后期空间紧凑、以单人体验为目标，并不是大型开放探索环境。它的作用是让玩家集中处理四段被封存的记忆及其对话，而不是提供大量导航内容。",
    ),
    30: (
        "The initial memory room is organised around repeated pick-up, inspection, listening, and placement. Its open circulation lets the player compare recent memories from a stable central position, while Focus, Context, and Background translate an abstract judgment into distance and location. After the body-interior sequence, the same room returns with a changed meaning: four released deep-memory objects appear at explicit return points and can again be handled and compared with the three recent objects. Reusing this room creates narrative contrast. A space that first supported a partial arrangement becomes the site of the player's broadest agency, where the consequences of the four container conversations are integrated through the final seven-object attention allocation.",
        "初始记忆空间围绕反复拿起、查看、聆听和放置来组织。开放的动线让玩家能够从稳定的中央位置比较近期记忆，而“焦点”“语境”和“背景”则把抽象判断转化为距离与位置。身体内部段落结束后，同一个房间以不同意义回归：四个释放出的深层记忆物件出现在明确的返回点，再次可以被拿取，并与三个近期物件一起比较。重复使用该房间形成叙事对照。最初只支持局部安排的空间，后来成为玩家能动性最强的场所，通过最终七物件注意力分配来整合四场容器对话产生的结果。",
    ),
    33: (
        "The body-interior chamber uses a compact single-player layout with four containers positioned as clear interaction destinations. After the forced swallow and fall, the player regains movement within this bounded space and can choose which container to approach first. Start Conversation appears only when the player enters the relevant proximity area, so limited movement still has a functional role: it links each location to one deep memory and its dialogue. When a memory gains a constructive interpretation, its container breaks and the object is released. After the fourth release, the space begins to fail and the story transfers the player back to the memory room. The spatial design is therefore intentionally focused and economical rather than expansive.",
        "身体内部空间采用紧凑的单人布局，四个容器作为清晰的交互目的地分布其中。经过强制吞咽与坠落后，玩家在这个边界明确的空间中重新获得移动能力，并可以选择先接近哪个容器。只有进入对应的接近区域后，“Start Conversation”才会出现，因此有限的移动仍具有功能：它把每个位置与一段深层记忆及其对话连接起来。当记忆获得建设性的新理解后，对应容器破碎并释放物件。第四段记忆释放后，空间开始失效，故事把玩家送回记忆空间。因此，这里的空间设计有意保持集中和经济，而不是追求拓展规模。",
    ),
    37: (
        "Sensory design distinguishes the office, memory room, Giant crisis, and body-interior stage through changes in lighting, post-processing, dialogue, and sound. The robot assistant provides a continuous narrative voice across these transitions. Its warning precedes the Giant's appearance, and its panicked reaction continues during the black-screen swallow, allowing sound to preserve continuity when the scene is visually hidden. Inside the body, the assistant's dialogue acknowledges the enclosed, resonant environment and later confirms when each container breaks and when all four memories have been released. These responses translate changes in danger and progress into character dialogue rather than relying only on interface text.",
        "感官设计通过灯光、后期处理、对白和声音变化，区分办公室、记忆空间、巨人危机与身体内部阶段。机器人助手的叙事声音跨越这些转场，保持体验连续。它在巨人出现前发出警告，并在黑屏吞咽期间继续以惊慌语气作出反应，使视觉被遮蔽时仍能通过声音维持连贯。进入身体内部后，助手对白会确认这个封闭、具有回响的环境，之后也会在容器破碎及四段记忆全部释放时确认进度。这些回应把危险与进度变化转译为角色对白，而不是只依赖界面文字。",
    ),
    41: (
        "Inside the body-interior space, echo is used to make the chamber feel larger than its compact layout. Reverb Zones and Audio Echo Filters add reverberation and delayed reflections to dialogue and environmental audio, creating the impression of a huge, hollow cave surrounding the player. The contrast between the visually bounded single-player space and its long acoustic tail increases immersion and suggests that the client's internal beliefs extend beyond what can be seen directly. The dialogue itself remains the main feedback system: constructive and harmful choices receive different spoken responses, while successful resolution is marked by the corresponding container breaking and releasing its memory.",
        "在身体内部空间中，回声被用来让紧凑布局产生更大的听觉尺度。Reverb Zone 与 Audio Echo Filter 为对白和环境声音加入混响及延迟反射，使玩家感觉自己处在一个巨大、空洞的洞穴中。视觉上边界明确的单人空间与较长的声音尾音形成对比，增强沉浸感，也暗示来访者的内部信念延伸到直接可见范围之外。对话本身仍然是主要反馈系统：建设性与有害选择会得到不同的话语回应，而成功解决则通过对应容器破碎并释放记忆来标记。",
    ),
    44: (
        "The crisis and body-interior soundscapes have different functions. Footsteps, heartbeat, roof rumble, groans, swallowing audio, black-screen dialogue, and the controlled fall make the climax feel immediate and bodily. After the climax, the falling-action chamber replaces that aggressive sequence with a resonant acoustic field created through reverb and echo processing. This shift allows the player to recognise that the story has moved from external threat to reflection and reinterpretation. Container-breaking sounds and the assistant's responses mark completed memories, while the final spatial breakdown prepares the return to the memory room. These sequences still require final end-to-end WebGL and XR timing checks, so the report describes the intended implemented setup without claiming completed platform validation.",
        "危机与身体内部阶段的声景承担不同功能。脚步、心跳、屋顶震动、呻吟、吞咽音效、黑屏对白与受控坠落，使高潮具有直接而具身的冲击。高潮之后，下降行动空间用通过混响与回声处理形成的共鸣声场取代这种攻击性序列。这样的变化让玩家意识到，故事已经从外部威胁转向反思与重新理解。容器破碎声和助手回应标记已经完成的记忆，最后的空间解体则为返回记忆空间作准备。这些序列仍需完成 WebGL 与 XR 的端到端时序检查，因此报告描述的是预期的已实现设置，而不声称平台验证已经完成。",
    ),
    48: (
        "The crystal ball and memory objects establish the project's first semiotic system: hidden mental content becomes visible, audible, and physically handled. The four later objects deepen this system because their ordinary forms carry specific beliefs about rank, rest, help-seeking, and mistakes. Before the conversations, the medal, alarm clock, old phone, and red correction pen are sealed inside four containers. Their confinement gives physical form to memories that the client has psychologically locked away. When released, the same objects remain intact and later return to the memory room, reinforcing that the Memory Organizer changes attention and interpretation rather than deleting the client's history.",
        "水晶球与记忆物件建立了项目的第一套符号系统：隐藏的心理内容变成可见、可听、可被身体操作的东西。后续四个物件进一步深化这套系统，因为它们的日常形态承载关于排名、休息、求助和犯错的具体信念。对话开始前，奖牌、闹钟、旧手机和红色批改笔被封存在四个容器中。它们被限制的状态，把来访者在心理上封锁起来的记忆转化为物理形式。释放后，同一批物件仍保持完整，并在之后回到记忆空间，再次强调记忆整理师改变的是注意力与解释，而不是删除来访者的历史。",
    ),
    49: (
        "The breaking containers are the central symbol of the body-interior stage. Each container breaks only after its memory receives a new constructive interpretation. This physical break represents the destruction of an old belief and the birth of a new understanding or worldview. At the same time, the memory changes from being sealed and restricted to being released and free. The image therefore also suggests that part of the client's psychological restraint has broken. The client gains freedom to reinterpret past experiences, to see the self from another position, and to decide that an old event does not have to impose the same meaning forever. The memory object survives the break, so liberation comes from changing its framework rather than erasing what happened.",
        "破碎的容器是身体内部阶段的核心符号。每个容器只有在其记忆获得新的建设性理解后才会破碎。这种物理破裂代表旧观念的瓦解，以及新理解或新世界观的诞生。与此同时，记忆也从被封存、被限制的状态转为被释放并恢复自由。因此，这个意象还暗示来访者心理上的一部分枷锁已经破碎。来访者由此获得重新解释过去经历、从不同位置看待自己，并决定旧事件不必永远施加同一种意义的自由。记忆物件在破碎过程中仍然存在，因此解放来自解释框架的改变，而不是抹除已经发生的事情。",
    ),
    50: (
        "The swallow transition combines blackness, swallowing audio, assistant dialogue, and a falling sequence into a readable sign of crossing into the Giant's body. The cavern-like reverb then establishes the body-interior chamber as a psychologically enlarged inner space even though its physical layout is compact. Container breakage and released objects signify movement from confinement toward reinterpretive freedom. Back in the memory room, Focus, Context, and Background extend the same non-moral semiotic logic. They represent degrees of present attention rather than good, neutral, and bad. The final report makes those symbolic positions concrete by describing how each allocation may affect the client's psychology.",
        "吞咽转场把黑暗、吞咽音效、助手对白和坠落组合成一个清晰的“进入巨人身体”符号。随后，洞穴般的混响把身体内部空间建立为心理上被放大的内在场所，尽管它的物理布局十分紧凑。容器破碎与物件释放象征记忆从受限状态走向重新解释的自由。回到记忆空间后，“焦点”“语境”和“背景”延续同一套非道德化符号逻辑。它们表示当下注意力的不同程度，而不是好、中、坏。最终报告通过描述每种分配可能如何影响来访者心理，使这些符号位置转化为具体后果。",
    ),
    54: (
        "Embodied action changes with each narrative stage. In the office, the player explores and activates the crystal ball; in the first memory room, they pick up objects, listen, and place them into attention zones. During the climax, movement is constrained while the Giant captures and swallows the player. In the falling-action body space, the player regains movement, approaches one of four containers, enters its interaction area, starts a conversation, and selects responses across two rounds. The player does not break a container through physical attack; the container breaks automatically when the memory gains a constructive interpretation. The final resolution returns to object handling, where the player exercises the widest embodied choice by redistributing all seven memories.",
        "具身动作会随叙事阶段改变。在办公室，玩家进行探索并激活水晶球；在第一个记忆空间中，玩家拿起物件、聆听内容，并把它们放入注意区域。高潮期间，巨人抓住并吞下玩家，移动受到限制。进入下降行动的身体内部空间后，玩家恢复移动，接近四个容器之一，进入其交互区域，开始对话，并在两轮中选择回应。玩家不会通过物理攻击打碎容器；当记忆获得建设性的新理解后，容器会自动破碎。最终解决阶段重新回到物件操作，玩家通过重新分配全部七段记忆，作出项目中最广泛的具身选择。",
    ),
    55: (
        "The capture and swallow create an enacted loss of control without switching to an external cinematic camera. The Giant tears away the roof, reaches into the room, aligns the player with the hand by moving the XR Origin, and carries them toward the animated mouth. Locomotion and jumping are disabled, but head direction or mouse look remains active, allowing the player to witness the change in scale and position. The black fade does not end the embodied sequence: swallowing audio continues, the player is moved to the fall start point, and ControlledFall delivers them to the body-interior landing point before locomotion is unlocked. This transition marks the movement from the climax into the falling-action container stage.",
        "抓取与吞咽在不切换到外部电影镜头的情况下，展演了控制权的丧失。巨人掀开屋顶、把手伸入房间，通过移动 XR Origin 让玩家与手部对齐，再把玩家带向动画中的嘴部。位移与跳跃被禁用，但头部方向或鼠标观察仍然有效，使玩家能够目睹尺度与位置的变化。黑色淡出并没有结束具身段落：吞咽音频继续，玩家被移动到坠落起点，ControlledFall 把玩家送到身体内部落点，之后才解除移动锁定。这个转场标志着故事从高潮进入下降行动的容器阶段。",
    ),
    56: (
        "The Giant, the assistant, the containers, and the environment all perform narrative actions. The Giant approaches, removes the roof, grabs, groans, and swallows, creating the climax. The assistant warns the player, reacts during the black transition, responds to progress in the body-interior space, and initiates the return after all four memories are released. Each container remains closed while its memory is unresolved and breaks after a constructive interpretation, releasing the object without destroying it. After the fourth release, the body-interior space shifts and breaks down. On return, the four freed objects reappear at deliberate positions in the familiar memory room, allowing the player to begin the final attention allocation.",
        "巨人、助手、容器与环境都在执行叙事动作。巨人接近、移除屋顶、抓取、呻吟并吞咽，从而形成高潮。助手向玩家发出警告，在黑屏转场中作出反应，回应身体内部空间中的进度，并在四段记忆全部释放后启动返回。每个容器会在记忆尚未解决时保持封闭，并在获得建设性新理解后破碎，在不摧毁物件的情况下将其释放。第四段释放后，身体内部空间移动并解体。返回时，四个恢复自由的物件在熟悉的记忆空间中预先设置的位置重新出现，使玩家可以开始最终注意力分配。",
    ),
}


STATE_TABLE = {
    (1, 0): ("Exposition", "开端"),
    (1, 1): ("OfficeDialogue", "OfficeDialogue"),
    (1, 2): (
        "Learn the role, explore, or begin work; these choices establish rhythm but reconverge.",
        "了解角色、探索或开始工作；这些选择改变节奏，但最终汇合。",
    ),
    (2, 0): ("Exposition to tension building", "开端至张力累积"),
    (2, 1): ("AwaitCrystalBall / TransitionToHippocampus", "AwaitCrystalBall / TransitionToHippocampus"),
    (2, 2): (
        "Finish the briefing and cross the threshold into the client's subconscious.",
        "完成简报并跨越门槛，进入来访者的潜意识。",
    ),
    (3, 0): ("Tension building", "张力累积"),
    (3, 1): ("AwaitMemoryPlacement", "AwaitMemoryPlacement"),
    (3, 2): (
        "Inspect and provisionally place three recent memories before the crisis interrupts.",
        "查看并暂时放置三段近期记忆，之后危机将其打断。",
    ),
    (4, 0): ("Climax", "高潮"),
    (4, 1): ("GiantCrisis / SwallowTransition", "GiantCrisis / SwallowTransition"),
    (4, 2): (
        "The Giant breaks the safe space, captures the player, and swallows them; looking remains available while movement is withheld.",
        "巨人破坏安全空间、抓住并吞下玩家；位移被暂停，但仍可观察。",
    ),
    (5, 0): ("Falling action", "下降行动"),
    (5, 1): ("Body-interior chamber / Four containers", "身体内部空间 / 四个容器"),
    (5, 2): (
        "Choose container order and dialogue responses; constructive interpretations break containers and release memories.",
        "选择容器顺序与对话回应；建设性新理解会使容器破碎并释放记忆。",
    ),
    (6, 0): ("Resolution", "解决"),
    (6, 1): ("FinalMemoryPlacement", "FinalMemoryPlacement"),
    (6, 2): (
        "Freely redistribute all seven memories; confirmed attention choices create real psychological outcomes.",
        "自由重新分配七段记忆；确认后的注意力选择产生真实心理后果。",
    ),
    (7, 0): ("Denouement", "结局"),
    (7, 1): ("BackToOffice / ViewReport / SessionComplete", "BackToOffice / ViewReport / SessionComplete"),
    (7, 2): (
        "Read the compositional report, then use Close Session for the neutral spoken ending and final fade to black.",
        "阅读组合式报告，再以 Close Session 触发中性口述结尾并最终淡为黑屏。",
    ),
}


def set_east_asia(run, font_name=CN_FONT):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def replace_bilingual_paragraph(paragraph, english, chinese, chinese_size=9.5):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    paragraph.add_run(english)
    paragraph.add_run().add_break()
    cn = paragraph.add_run(chinese)
    set_east_asia(cn)
    cn.font.size = Pt(chinese_size)
    cn.font.color.rgb = CN_GREY


def replace_bilingual_cell(cell, english, chinese):
    replace_bilingual_paragraph(cell.paragraphs[0], english, chinese, chinese_size=8.2)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def main():
    doc = Document(SOURCE)
    if len(doc.paragraphs) != 92 or len(doc.tables) != 4:
        raise RuntimeError(
            f"Unexpected source structure: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
        )

    for index, (english, chinese) in REVISIONS.items():
        replace_bilingual_paragraph(doc.paragraphs[index], english, chinese)
        doc.paragraphs[index].paragraph_format.keep_together = True

    state_table = doc.tables[2]
    for (row_index, cell_index), (english, chinese) in STATE_TABLE.items():
        replace_bilingual_cell(state_table.rows[row_index].cells[cell_index], english, chinese)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
