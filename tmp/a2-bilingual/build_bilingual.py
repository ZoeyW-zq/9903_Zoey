from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\GitHub\9903_Zoey")
SOURCE = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer.docx"
OUTPUT = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer_Bilingual.docx"
FLOWCHART = Path(r"E:\Desktop\9903NarrativeAndSensemaking\A2\freytag_flowchart.png")
ENDINGS = ROOT / "tmp" / "a2-bilingual" / "actual_endings_table.png"

CN_FONT = "Microsoft YaHei"
CN_GREY = RGBColor(0x58, 0x62, 0x70)
PLACEHOLDER_RED = RGBColor(0xC0, 0x00, 0x00)


HEADINGS = {
    0: "DDES9903 - Assessment 2 报告草稿",
    1: "记忆整理师",
    3: "摘要",
    6: "当前实现与开发进度",
    10: "叙事状态网络",
    15: "玩家能动性",
    21: "叙事连贯性",
    26: "空间的运用",
    30: "感官触发与多模态元素的运用",
    35: "符号元素的运用",
    39: "具身动作与展演元素的运用",
    44: "协作实践",
    47: "学术参考文献",
    55: "资产引用",
    63: "AI 使用说明",
}


PARAGRAPHS = {
    5: "《记忆整理师》是一段关于“注意力如何改变记忆的意义与情绪力量”的短篇互动叙事。玩家与一名桌面助手合作，进入一位办公室职员的潜意识，在不删除或改写任何记忆的前提下整理七段记忆。体验从办公室任务简报开始，随后进入记忆空间；由羞耻、焦虑和压力累积而成的巨型小丑引发危机，玩家在其内部完成四场对话式对抗，最后返回办公室。玩家能动性体现在入职引导节奏、镜子对话的处理顺序与回应路径，以及把记忆实体重新分配到“焦点”“语境”和“背景”三个注意区域的具身行动中。最终报告把这些位置解释为定性的注意模式，而不是分数或唯一正确结局。这样的设计在保持完整戏剧弧线的同时，也允许玩家表达自己对来访者当前需要的理解。",
    7: "截至当前版本，scene_WebGL.unity 中从办公室到最终组合式报告的五阶段叙事主线已经实现。2026 年 8 月 16 日的脚本构建以零错误完成。Git 工作记录也显示了办公室分支、海马体对话、记忆分配、巨人内部场景、音频和最终报告的分阶段开发。报告系统保存七个最终位置，并为每段记忆从三个已写好的文本片段中选择一个，因此只需 21 个片段即可支持 2,187 种最终配置。这是一种有意识的制作策略，并不意味着项目手工撰写了 2,187 个完整故事。",
    9: "[提交前补充：实测单次游玩时长。模板建议单次体验最长不超过五分钟。]",
    12: "图 1：按 Freytag 戏剧阶段标注的叙事流程图。",
    13: "图 1 将运行时状态机映射到 Freytag 戏剧弧线，并区分必然发生的戏剧进程与局部能动性。办公室选项在进入潜意识前重新汇合。无论初步分配如何，放置表层记忆都会触发危机，因为巨人源于被压抑的痛苦记忆，而不是对“错误答案”的惩罚。镜室是一个枢纽：四段记忆可以按任意顺序处理，有害回应会形成局部循环，但不会破坏后续状态。第四场成功对话构成戏剧高潮。随后，玩家在下降行动阶段重新分配全部七段记忆，获得最强的表达性选择。结局不是一个预写好的固定终点，而是根据确认后的注意模式组合而成的报告。",
    16: "《记忆整理师》的能动性是分层分布的，而不是在每一时刻保持相同强度。在办公室阶段，玩家控制入职引导的节奏：可以先了解工作内容、探索办公室、返回，再开始工作；已使用的工作说明选项会消失，而探索可以重复。这些选择形成至少五条有意设计的入职路径，但最终都汇合到同一位来访者的任务，从而在提供好奇和准备空间的同时，不成倍增加后续制作量。",
    17: "最强的手工分支出现在巨人内部。四面镜子可按任意顺序处理，仅顺序就产生 24 种排列，尚未计入对话差异。每面镜子包含两轮对话。第一轮中的两个合理回应会进入不同的第二轮问题；任一轮中的有害捷径都会结束当前尝试、强化局部声音和环境，再恢复“开始对话”。失败会改变玩家当下的体验，但不会永久移除内容，也不会制造不连贯的全局状态。错误因此被叙事化为内部信念的持续存在，而不是传统的游戏结束画面。",
    18: "最终重新分配是主要的表达性能动性。每段记忆都可以在“焦点”“语境”和“背景”之间移动，并可在确认前反复修改。七个相互独立的三元选择产生 3^7，即 2,187 种注意配置。这些是真实的状态差异：报告会显示每段记忆的最终类别，并为每个物件选择对应的结果片段。但系统避免了简单的好坏计分；“焦点”不一定是正面结果，“背景”也不等于删除。玩家决定的是哪些内容应占据当前注意力，而不是改写来访者的人生。",
    19: "该方法有意结合真实选择与选择幻觉。第三段表层记忆放置后，巨人危机必然发生，因此初步布局不会绕开核心冲突。在系统层面这是汇流；在故事内部，危机则来自痛苦记忆长期被排除。自由的查看顺序、空间位置修改、局部对话循环，以及真正改变报告的最终状态，共同维持了故事世界看似开放的感受。其创新点在于“组合式后果”：少量经过控制的文本单元能够产生大量连贯解释，而无需假装模拟一个没有边界的故事。",
    20: "仍有一项限制：在抓取和吞咽段落中，系统会暂时锁定平移移动，避免玩家逃离已编排的过场，但不会强制覆盖相机旋转或头部追踪方向。这是对导航的短暂限制，并非脱离玩家视角的屏幕式过场。最终报告应如实说明这一点，而不应声称移动从未中断。",
    22: "设计以一条固定的主题规则开始：记忆整理师只能重新分配注意力，不能删除、改写或制造记忆。这条规则保护了所有路径上的因果关系。来访者已经内化了一个信念：只有表现完美、承受压力并且不给别人添麻烦，自己才有价值。三段近期记忆和四段深层记忆从不同角度检验这一信念。水瓶重复了带病工作；夕阳和积木提供了其他价值来源；奖牌、闹钟、旧手机和红笔分别具象化完美主义、休息罪恶感、孤立以及对犯错的恐惧。由于所有物件都连接到同一信念系统，不同查看顺序仍然保持连贯。",
    23: "项目通过把固定戏剧主线与可变内容分开来维持 Freytag 结构。开端与触发性转场建立角色和规则；表层记忆放置逐步提出“缺失了什么”的问题；巨人的出现升级冲突，四场镜子对话迫使来访者面对支撑巨人的信念；第四段记忆被化解时，巨人失去支撑，形成高潮。最终重新分配属于下降行动，玩家把新理解应用到全部记忆，而不是面对第二个敌人。办公室报告通过反映玩家选择的注意模式完成结局。",
    24: "多个门控机制防止顺序错乱。只有在“开始工作”对话和来访者报告阅读完成后，水晶球才会出现；只有三个不同的表层物件都被放置后，初始危机才会开始；每面镜子只有在第二轮选择建设性回应后才计入一次进度；最终确认需要七段不同记忆全部被放置；报告只读取已确认的最终位置字典。这些限制明确了状态转移，同时保留了状态内部的顺序自由与解释空间。",
    25: "设计也控制了未来制作负担。四面镜子共享同一个数据驱动的 MemoryDialogueController，而不是四套独立脚本；初始分配与最终分配复用同一个分阶段 MemoryPlacementController。最重要的是，报告使用 21 个物件级片段，而不是 2,187 份完整报告。每个片段只描述一段记忆在一个注意类别中的含义，因此组合后不会与某个全局“成功”标签冲突，因为系统根本不设置这种标签。这既是实现层面的节省，也是拒绝把心理健康简化成分数的叙事选择。",
    27: "空间首先承载叙事意义，其次才承担导航功能。普通办公室把记忆工作框定为一项专业任务，也为助手提供了一个叙事世界内部的简报地点。水晶球是行政现实与来访者潜意识之间的门槛物件。记忆空间把注意力外化：“焦点”“语境”和“背景”不是菜单选项，而是物理区域。移动物件会改变它与玩家身体及其他记忆的关系，使抽象的认知判断成为可在空间中检查的东西。",
    28: "故事随后改变尺度。巨型小丑破坏记忆空间的表面安全感、掀开屋顶并吞下玩家。在巨人内部，四面镜子把痛苦信念彼此分开，同时又把它们封闭在同一个身体里。玩家可以按任意顺序接近镜子，因此镜室是叙事枢纽，而非单向走廊。当四段记忆都被化解后，空间解体，释放出的物件回到受损的记忆空间。破损的天花板让危机留下持续的空间痕迹，即使剧情已经进入下降行动。",
    29: "在叙事连贯性需要的地方，空间能动性会受到边界约束。放置区域允许反复调整，而不重叠的触发体和不同的物件身份确保每段最终记忆只有一个可读状态。状态机按阶段启用办公室、海马体和噩梦根节点，使 XR 中只有当前相关的故事空间处于激活状态，从而降低运行成本。当前版本仍需在 Play Mode 中检查区域触发重叠、直接跨区移动、返回点对齐，以及目标 WebGL 分辨率下的可读性。",
    31: "音频、灯光、运动和淡入淡出被用作状态信息，而不只是装饰。拿起记忆会播放与其相关的叙述，把物件操作与声音连接起来。第二个表层物件放置后，低沉声音、灯光波动和房间扰动共同提示可见记忆并不完整。危机期间，脚步声结束后才出现隆隆声，巨人逐步接近，屋顶和玩家都被移动到同一个手部锚点。这些线索把抽象的压力积累转化为正在逼近的物理威胁。",
    32: "不同转场使用不同的视觉语言。白色淡出伴随水晶球把玩家送入潜意识，暗示连接与进入；黑色淡出与移向巨人口部的最后一段运动重叠，并掩护传送到吞咽/坠落空间的不连续性。Global Volume 权重区分办公室、海马体和噩梦状态；四面镜子被化解后，噩梦视觉处理逐渐淡出。危机开始时粒子发射量也会下降，让环境本身传达系统状态改变，而不只依赖字幕。",
    33: "镜子音频创造持续存在感。每面未解决的镜子都会循环播放负面内部陈述。有害回应会强化声音与周围环境，然后让镜子回到待机状态；建设性解决会停止回声并释放原始物件。因此，玩家能够听见整个房间随着进度逐渐安静。在 VR 中，持续触碰水晶球还能提供手柄触觉反馈；WebGL 版本则以点击交互替代。这种平台差异应由不同证据分别展示，而不应描述成完全相同。",
    34: "多模态线索由控制叙事进度的同一组状态控制器排序，从而降低相互矛盾的风险。剩余验证仍然重要：循环播放的脚步声可能阻止危机继续；黑屏转场期间的字幕与音频时序仍需检查；最长的报告布局也必须验证可读性。这些是当前风险，而不是已完成的测试结果。",
    36: "核心符号系统是记忆物件与注意力之间的关系。物件刻意保持日常：用过的水瓶、夕阳照片、积木、第二名奖牌、闹钟、旧手机和红色批改笔。熟悉感让个人意义从语境中产生。水瓶象征以健康换取坚持；闹钟把休息变成罪恶感；手机上未发送的求助信息象征害怕成为负担；红笔把一次错误变成对能力的总判决。奖牌与积木对比外部排名和耐心完成后的私人满足，夕阳则代表不需要生产力或他人认可也可以成立的价值。",
    37: "巨型小丑把来访者记忆中已有的小丑形象，与长期累积到难以控制的焦虑、羞耻、压力和自我批评结合起来。它不是随机怪物，也不是错误选择的惩罚。镜子象征重复的内部语言：它们把来访者的信念反射成看似客观的事实。成功对话后镜子自动破裂，因为任务是重新理解，而不是实施物理攻击。原始物件仍然存在，再次强化“记忆没有被抹除”的规则。",
    38: "“焦点”“语境”和“背景”避免使用道德化标签。它们描述的是注意距离，最终报告只使用定性的彩色类别文字，不使用成败色、星级或百分比。同一位置可能具有混合含义：把痛苦记忆放到背景可以减轻即时负担，但也可能让其经验教训更难被调用；把积极记忆放在焦点可以带来修复，但系统不会宣布一个普遍正确的最优解。因此，该符号系统在支持能动性的同时，也保护了“痛苦与愉快经历会继续共存”的伦理主张。",
    40: "玩家通过对隐喻采取行动来理解叙事。他们探索办公室、激活水晶球、拿起记忆物件、在持有时聆听内容、把物件放进注意区域、接近镜子、开始对话，并在确认前修改最终布局。这些行为使认知过程变得可触：注意力成为距离和位置；对抗成为靠近反射表面；重新理解会释放一个之后仍可再次拿起的物件。",
    41: "具身体验也改变了戏剧风险。巨人掀开屋顶、伸手抓住玩家、把玩家带向口部，并触发吞咽/坠落转场。系统只按位置移动 XR Origin，同时保留头部追踪方向，因此玩家仍然感知自己处于场景内部，而不是观看外部化身。VR 与 WebGL 共享移动锁定：抓取期间停止位移，但仍允许鼠标观察或头部转动。这种被编排的能动性丧失支撑危机段落，随后玩家在镜室重新获得控制，并在最终分配中拥有最强的表达自由。",
    42: "助手也以行动推动叙事。它向玩家简报、注意到缺失的记忆、对危机作出反应，在四段回声停止后重新获得传送信号，并把玩家和释放出的物件带回记忆空间。由于这些行动与完成事件绑定，助手是故事世界中的参与者，而不是漂浮的说明面板。叙事内的电脑页面、世界空间选项、镜子、按钮和物理区域也进一步减少了对屏幕固定 UI 的依赖。",
    43: "[补充：XR 头显测试证据、设备名称、舒适度观察、帧率结果和一张带标签截图。缺少这些证据时，不要声称已满足 c7。]",
    45: "[此部分留给学生整理。插入第 6-10 周官方课程频道中的英文截图。每项证据都应标注准确的评分标签（d1-d12）、日期、频道、你贡献了什么，以及该贡献带来了什么变化。]",
    54: "[需要核对：按照课程要求统一引用格式，并补充实际使用的第 2-4 周和第 8 周课程材料。]",
    56: "以下来源可以在仓库文档中确认。最终提交只应保留实际出现在构建中的资产，并为每个导入的模型、图像、声音、字体和动画提供准确来源与许可。",
    62: "[需要核对并引用：Whiteclown N Hallin 模型/动画、Ace Combat 7 Medals ACE 模型、闹钟、手机、可破碎玻璃、BrickToy 资产、所有记忆图像，以及所有助手、记忆、环境、心跳、隆隆声、脚步、怪物和吞咽音频。删除最终版本未使用的继承示例引用。]",
    64: "OpenAI Codex 被用作开发与写作助手。在本报告草稿中，Codex 阅读了项目上下文和评估文档，对照详细评分表检查报告结构，生成初始叙事状态网络图，并起草供学生后续审核的文本。AI 输出被当作草稿处理：项目陈述只采用仓库中已有的证据；未知链接、协作证据、资产来源、时长和头显测试结果都保留为占位符。",
}


TABLES = {
    0: [["草稿状态：完成所有高亮占位符，确认五分钟游玩目标，核对全部资产来源，并在提交前删除本提示。"]],
    1: [
        ["提交字段", "草稿值"],
        ["WebGL 链接", "[待补充 - 可公开游玩的链接]"],
        ["GitHub 仓库", "GitHub 仓库链接"],
        ["首个场景", "scene_WebGL"],
        ["学生姓名 / zID", "[待补充]"],
    ],
    2: [
        ["当前项目证据", "状态及其对报告的意义"],
        ["办公室、表层记忆、镜子、最终分配和报告状态", "已实现并绑定到场景，支持完整叙事论证。"],
        ["分支对话与四场镜子对话", "已实现；全部镜子顺序与失败/重试路径仍需全流程测试。"],
        ["WebGL 端到端运行与世界内报告可读性", "仍待验证；在取得证据前不要描述为已完全测试。"],
        ["XR 头显部署", "XR 架构已存在，但给定项目上下文中没有部署证据。"],
    ],
    3: [
        ["Freytag 功能", "叙事状态", "该状态承载的能动性"],
        ["开端", "OfficeDialogue", "了解角色、重复探索或开始工作。"],
        ["触发事件", "AwaitCrystalBall / TransitionToHippocampus", "决定何时结束阅读并激活转场。"],
        ["上升行动", "AwaitMemoryPlacement", "查看并初步放置三段表层记忆。"],
        ["冲突升级", "GiantCrisis / SwallowTransition", "空间控制被具身化地剥夺，但头部视角仍可控制。"],
        ["高潮", "MirrorChamber", "选择镜子顺序与对话回应；失败会形成重试循环。"],
        ["下降行动", "FinalMemoryPlacement", "修改七段记忆各自的焦点/语境/背景分配。"],
        ["结局", "BackToOffice / View Report", "阅读个性化后果模式并返回电脑首页。"],
    ],
    4: [
        ["评分标签", "日期与频道", "证据截图", "该贡献为何具有实质性 / 互惠性"],
        ["[d__]", "[待补充]", "[插入英文截图]", "[待补充]"],
        ["[d__]", "[待补充]", "[插入英文截图]", "[待补充]"],
        ["[d__]", "[待补充]", "[插入英文截图]", "[待补充]"],
    ],
    5: [
        ["服务", "用途", "需要保留的提示词 / 记录"],
        ["OpenAI Codex", "报告结构、初稿、图表和文档排版", "2026 年 8 月 16 日用户提示：阅读 PROJECT_CONTEXT 和三份评估文件，起草报告，协作部分稍后补充。"],
        ["OpenAI Codex", "Unity 实现、调试、测试或场景分析", "[粘贴相关的准确提示词，或提供链接/附录记录]"],
    ],
}


def set_east_asia(run, font_name=CN_FONT):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def make_bilingual_heading(paragraph, chinese):
    english = paragraph.text
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run(english + " / ")
    cn = paragraph.add_run(chinese)
    set_east_asia(cn)


def append_chinese(paragraph, chinese, placeholder=False, size=9.5):
    br = paragraph.add_run()
    br.add_break()
    cn = paragraph.add_run(chinese)
    set_east_asia(cn)
    cn.font.size = Pt(size)
    cn.font.color.rgb = PLACEHOLDER_RED if placeholder else CN_GREY
    if placeholder:
        cn.bold = True
        cn.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_alt_text(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def clone_sectpr(base_sectpr, landscape=False):
    sect_pr = deepcopy(base_sectpr)
    type_node = sect_pr.find(qn("w:type"))
    if type_node is None:
        type_node = OxmlElement("w:type")
        sect_pr.insert(0, type_node)
    type_node.set(qn("w:val"), "nextPage")
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if landscape:
        width = pg_sz.get(qn("w:w"))
        height = pg_sz.get(qn("w:h"))
        pg_sz.set(qn("w:w"), height)
        pg_sz.set(qn("w:h"), width)
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.attrib.pop(qn("w:orient"), None)
    return sect_pr


def attach_section_break(paragraph, sect_pr):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(sect_pr)


def insert_endings_figure(doc, target_paragraph):
    base_sectpr = doc._element.body.sectPr

    portrait_break = doc.add_paragraph()
    attach_section_break(portrait_break, clone_sectpr(base_sectpr, landscape=False))

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    shape = image_paragraph.add_run().add_picture(str(ENDINGS), width=Cm(24.0))
    add_alt_text(
        shape,
        "Implemented Final Report Outcomes",
        "A bilingual table of the 21 one-sentence outcomes implemented for seven memory objects across Background, Context, and Focus attention.",
    )

    caption = doc.add_paragraph(
        "Figure 2. The 21 one-sentence outcomes implemented in the final report UI. / "
        "图 2：最终报告 UI 中实际实现的 21 条单句结局。",
        style="Figure Caption",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    for run in caption.runs:
        set_east_asia(run)

    landscape_break = doc.add_paragraph()
    attach_section_break(landscape_break, clone_sectpr(base_sectpr, landscape=True))

    for element in [portrait_break._p, image_paragraph._p, caption._p, landscape_break._p]:
        target_paragraph._p.addprevious(element)


def replace_flowchart(doc):
    drawing_paragraph = next(p for p in doc.paragraphs if p._p.xpath(".//w:drawing"))
    for run in list(drawing_paragraph.runs):
        drawing_paragraph._p.remove(run._r)
    drawing_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drawing_paragraph.paragraph_format.page_break_before = True
    drawing_paragraph.paragraph_format.keep_with_next = True
    # The supplied artwork is unusually tall; 9.5 cm keeps the full image and caption
    # together within the printable height of one A4 portrait page.
    shape = drawing_paragraph.add_run().add_picture(str(FLOWCHART), width=Cm(9.5))
    add_alt_text(
        shape,
        "The Memory Organizer Freytag Flowchart",
        "A six-stage flowchart showing exposition, tension building, crisis, deep memories, resolution, and denouement.",
    )


def update_footer(doc):
    seen = set()
    for section in doc.sections:
        footer = section.footer
        if id(footer._element) in seen:
            continue
        seen.add(id(footer._element))
        for paragraph in footer.paragraphs:
            if "DDES9903 A2 Report Draft" not in paragraph.text:
                continue

            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)

            label = paragraph.add_run("DDES9903 A2 Bilingual Draft / 中英对照工作稿  |  ")
            set_east_asia(label)

            page_field = OxmlElement("w:fldSimple")
            page_field.set(qn("w:instr"), "PAGE")
            field_run = OxmlElement("w:r")
            field_text = OxmlElement("w:t")
            field_text.text = "1"
            field_run.append(field_text)
            page_field.append(field_run)
            paragraph._p.append(page_field)

    for section in doc.sections:
        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        if page_numbering is not None:
            page_numbering.attrib.pop(qn("w:start"), None)


def move_ai_table_to_new_page(doc):
    table = doc.tables[5]
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    table._tbl.addprevious(page_break._p)


def main():
    doc = Document(SOURCE)

    for index, chinese in HEADINGS.items():
        make_bilingual_heading(doc.paragraphs[index], chinese)

    placeholder_indices = {9, 43, 45, 54, 62}
    for index, chinese in PARAGRAPHS.items():
        append_chinese(doc.paragraphs[index], chinese, placeholder=index in placeholder_indices)

    for table_index, translations in TABLES.items():
        table = doc.tables[table_index]
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if row_index >= len(translations) or cell_index >= len(translations[row_index]):
                    continue
                text = translations[row_index][cell_index]
                paragraph = cell.paragraphs[-1]
                placeholder = "待补充" in text or "插入英文截图" in text or "粘贴相关" in text
                append_chinese(paragraph, text, placeholder=placeholder, size=8.2)

    replace_flowchart(doc)
    compositional_agency = next(
        p for p in doc.paragraphs if p.text.startswith(
            "This approach deliberately combines real and illusory choice."
        )
    )
    insert_endings_figure(doc, compositional_agency)
    update_footer(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
