from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

import build_bilingual as base


A1_REPORT = Path(r"E:\Desktop\9903NarrativeAndSensemaking\A1\DDES9903 A1 Report_Zoey.docx")
MEDIA_DIR = base.ROOT / "tmp" / "a2-bilingual" / "a1-media"


ENGLISH_REVISIONS = {
    5: (
        "Memory Organizer is a short immersive interactive story about repairing an anxious "
        "Dreamer's subconscious. The player takes the role of a newly hired Memory Organizer "
        "on their first assignment. In the office, a robot assistant introduces the job and "
        "shows the sleeping Dreamer through a crystal ball. After entering the Dreamer's "
        "subconscious, the player organizes recent memories in a hippocampus-inspired memory "
        "space. With the assistant's guidance, the player reviews three recent memories and "
        "places them in Focus, Context, or Background attention zones. Just as the task seems "
        "under control, the dream world breaks apart and a giant clown appears as a manifestation "
        "of the Dreamer's anxiety and internalized pressure. The player is swallowed and enters "
        "a space formed by painful memories inside the giant clown. Four mirror conversations "
        "help the Dreamer face the beliefs connected to those memories. After returning to the "
        "memory room, the player can rearrange all seven memories across the three attention "
        "zones. The final report reflects those choices, and the mission ends with the player "
        "returning to the office."
    ),
    13: (
        "Figure 1 shows how the story follows Freytag's model while still allowing player "
        "choice. The office choices change how quickly the player begins the assignment, but "
        "they return to the same main task. The three recent memories lead to the Giant crisis. "
        "Inside the Giant, the four mirrors can be visited in any order, and an unhelpful answer "
        "returns that mirror to the start of its conversation. After all four memories are "
        "released, the player returns to the memory room and rearranges all seven memories. "
        "The story therefore keeps a clear exposition, crisis, climax, falling action, and "
        "denouement while allowing choices inside each stage."
    ),
    16: (
        "The office gives the player a small but clear form of agency. The player can ask about "
        "the job, explore the room, return to the assistant, and decide when to begin. These "
        "options create different first-day rhythms without changing the client or removing the "
        "main assignment. This keeps the introduction manageable while still making the player "
        "feel like a newcomer choosing how to prepare."
    ),
    17: (
        "The strongest branching happens inside the giant clown. The four mirrors can be visited "
        "in any order, which already creates 24 possible orders. Each mirror uses a two-round "
        "conversation. A constructive answer moves the conversation forward, while a harmful "
        "shortcut strengthens the negative voice and returns the mirror to the start. The player "
        "can try again, so failure becomes part of the Dreamer's continuing self-doubt rather "
        "than a game-over screen."
    ),
    18: (
        "The final memory arrangement gives the player the widest choice. Each of the seven "
        "memories can be placed in Focus, Context, or Background, and the player can change the "
        "layout before confirming it. This creates 3^7, or 2,187, possible attention patterns. "
        "The report changes for every memory according to its final zone. Focus is not always "
        "good and Background does not delete a memory; the player is deciding what the Dreamer "
        "should pay attention to now."
    ),
    19: (
        "Some choices are deliberately limited. The Giant crisis always begins after the third "
        "recent memory is placed, because the painful memories must still be confronted for the "
        "story to reach its climax. The player cannot avoid this event, but can choose the order "
        "of inspection, revise placements, retry mirror conversations, and create a genuinely "
        "different final report. The experience therefore combines a fixed dramatic structure "
        "with choices that change how the player reaches and interprets the ending."
    ),
    20: (
        "During the scripted capture, the player's movement is temporarily constrained, but the "
        "user can still freely look around in 3D. This reduces agency, but it fits the narrative "
        "and physical logic: if a giant has caught and swallowed the player, walking away freely "
        "should not be possible. Reduced agency therefore becomes part of the embodied experience "
        "of fear, pressure, and loss of control."
    ),
    22: (
        "The main story rule is that a Memory Organizer can move attention but cannot delete, "
        "rewrite, or invent memories. The seven objects all connect to the Dreamer's fear that "
        "personal value depends on perfect work, endurance, and not burdening other people. The "
        "water bottle, sunset photograph, LEGO bricks, medal, alarm clock, old phone, and red "
        "correction pen show different sides of this belief. Because they belong to the same "
        "emotional problem, the story remains coherent even when the player sees them in a "
        "different order."
    ),
    23: (
        "The structure follows Freytag's model while adapting it for an interactive experience. "
        "The player first understands the job, then performs a manageable memory task, and then "
        "loses control during the clown interruption. The four mirror conversations form the "
        "climax because they remove the beliefs supporting the Giant. The final arrangement is "
        "the falling action, where the player applies what they have learned. The office report "
        "provides the denouement by showing the result of the player's choices."
    ),
    24: (
        "Several gates keep the sequence clear. The crystal ball appears only after the player "
        "finishes the office briefing and reads the client report. The crisis begins only after "
        "three different recent memories have been placed. A mirror counts as complete only after "
        "a constructive second answer, and the final report can be confirmed only after all seven "
        "memories have a zone. These rules prevent broken story states while leaving freedom inside "
        "each stage."
    ),
    25: (
        "The same structure also keeps the project practical. All four mirrors use one dialogue "
        "controller, and the early and final placement stages reuse one memory placement system. "
        "The final report uses three written outcomes for each of seven objects, so 21 short text "
        "fragments can support 2,187 combinations. This allows many endings without writing a "
        "separate complete story for every arrangement."
    ),
    27: (
        "The office is arranged as a warm client-facing care space rather than a clinical "
        "laboratory. Ryan's distinction between space and place frames this location not as a "
        "neutral container, but as an environment with emotional and narrative meaning (Ryan, "
        "n.d.). The central table gives the player an obvious orientation point, while the sofa, "
        "entertainment objects, ornaments, and domestic details make the room feel less clinical. "
        "These objects support the fiction that this is a place for care, conversation, and "
        "gentle memory repair."
    ),
    28: (
        "The hippocampus memory room is designed as a room assembled from objects strongly "
        "remembered by the client. For this reason, the items in the space are not arranged with "
        "strict visual regularity; they feel like fragments from everyday life that have been "
        "kept in memory. Placing the table in the middle guides the player's body and attention: "
        "the player moves toward the table, picks up a memory object, listens to its content, and "
        "then places it in Focus, Context, or Background. In the final stage, the same room allows "
        "all seven objects to be compared and rearranged."
    ),
    29: (
        "When the clown appears, the environment expands into a hostile exterior with red fog, "
        "dark sky, and unsafe scale. The giant hand tearing open the roof makes the player feel "
        "small and powerless, which supports the story moment where anxiety overwhelms the ordered "
        "memory system. Inside the clown, stone walls, torches, chains, bones, and four mirror "
        "structures create a trapped, heavy atmosphere. The mirrors can be approached in any "
        "order, so this space works as a hub rather than a single corridor. After they are resolved, "
        "the player returns to the damaged memory room, where the broken roof remains as evidence "
        "of the crisis."
    ),
    31: (
        "The sensory design uses three distinct Global Volumes for different narrative stages, "
        "while the robot assistant's voice connects the whole experience. The assistant explains "
        "the task, reduces confusion during transitions, and keeps the player aware of the plot "
        "when the visuals become unstable."
    ),
    32: (
        "In the office-to-memory transition, the player places a hand on the crystal ball. This "
        "triggers continuous haptic vibration, then the view gradually blurs and fades to white. "
        "The warm office atmosphere gives way to a sensory threshold, making entry into the "
        "client's memory space feel deliberate and believable. In the WebGL version, a click "
        "starts the same transition without controller haptics."
    ),
    33: (
        "In the memory stage, a separate Global Volume gives the subconscious a blue-purple tone, "
        "floating particles, and stronger bloom. The particles make the air feel dreamlike, while "
        "spot lights emphasize the memory objects and attention zones. Picking up an object also "
        "plays its memory narration, so image, sound, and physical handling direct attention to "
        "the same story fragment."
    ),
    34: (
        "In the crisis stage, the nightmare Global Volume lowers exposure and adds a red vignette. "
        "Tension is intensified by layered audio: the giant's footsteps, strong heartbeat, tinnitus, "
        "the roof being lifted away, and dark background ambience. After the giant clown swallows "
        "the player, the screen turns black, but swallowing sounds and the assistant's voice continue "
        "so the player still understands the current situation. Inside the Giant, each unresolved "
        "mirror repeats a negative statement; the room becomes quieter as the player resolves them."
    ),
    36: (
        "Crystal ball: this fantasy-like object makes the transition into the subconscious feel "
        "natural. Instead of using a purely technical interface, the crystal ball suggests vision, "
        "diagnosis, and access to hidden mental content. Memory objects: the physical objects on "
        "the table turn invisible mental events into things the player can approach and handle. "
        "They make memory organization feel embodied rather than only explained through dialogue."
    ),
    37: (
        "Memory canvas: when the player picks up an object, the memory canvas shows the specific "
        "memory fragment. The image and narration connect an everyday object to the Dreamer's "
        "experience. Giant clown: the clown represents the Dreamer's negative emotions, internalized "
        "pressure, and fear of judgment. Its exaggerated scale makes anxiety visible as something "
        "that has grown too large and now interrupts the repair work. The mirrors repeat the "
        "Dreamer's harmful beliefs and break only after those beliefs are reinterpreted."
    ),
    38: (
        "After the player is swallowed, the assistant's voice, swallowing sound effects, and the "
        "falling-through-a-pipe image work together as signs that communicate a clear new situation: "
        "I have been swallowed, and I am now inside the giant's body. Focus, Context, and Background "
        "also work as signs. They describe how close a memory is to current attention without calling "
        "it good or bad. A painful memory can move to Background without being erased, while a positive "
        "memory can move to Focus without becoming the only correct answer."
    ),
    40: (
        "The embodied actions focus on two agents: the player and the giant clown. For the player, "
        "the key actions are touching the crystal ball, handling memory objects, placing them in "
        "attention zones, being captured by the giant clown, approaching mirrors, choosing dialogue "
        "responses, and rearranging all seven memories. These actions use embodied metaphors, where "
        "physical interaction helps the player understand an abstract process through the body "
        "(Antle et al., 2009). Attention becomes distance and placement, while confronting a harmful "
        "belief becomes approaching and speaking to a mirror."
    ),
    41: (
        "During the scripted capture, the player's movement is temporarily constrained, but the "
        "user can still freely look around in 3D. This reduces agency, but it fits the narrative and "
        "physical logic: if a giant has caught and swallowed the player, walking away freely should "
        "not be possible. Reduced agency therefore becomes part of the embodied experience of fear, "
        "pressure, and loss of control."
    ),
    42: (
        "The giant clown's actions form the crisis. It tears open the roof, reaches into the safe "
        "workspace, grabs the player, and swallows them. These actions drive the plot while creating "
        "strong visual and auditory shock. The robot assistant also acts inside the story: it briefs "
        "the player, reacts to the crisis, restores the return signal after all four mirrors are "
        "resolved, and brings the player and released objects back to the memory room."
    ),
    48: "Antle, A. N., Corness, G., & Droumeva, M. (2009). What the body knows: Exploring the benefits of embodied metaphors in hybrid physical digital environments. Interacting with Computers, 21(1-2), 66-75. https://doi.org/10.1016/j.intcom.2008.10.005",
    49: "Cunliffe, A. L., & Coupland, C. (2012). From hero to villain to hero: Making experience sensible through embodied narrative sensemaking. Human Relations, 65(1), 63-88. https://doi.org/10.1177/0018726711424321",
    50: "Freytag, G. (1894). Freytag's Technique of the Drama (E. J. MacEwan, Trans.). Scott, Foresman and Company. (Original work published 1863).",
    51: "Gödde, M., Gabler, F., Siegmund, D., & Braun, A. (2018). Cinematic narration in VR: Rethinking film conventions for 360 degrees. In Virtual, Augmented and Mixed Reality: Applications in Health, Cultural Heritage, and Industry (pp. 184-201). Springer. https://doi.org/10.1007/978-3-319-91584-5_15",
    52: "Jenkins, H. (2004). Game design as narrative architecture. In N. Wardrip-Fruin and P. Harrigan (Eds.), First Person: New Media as Story, Performance, and Game (pp. 118-130). MIT Press.",
    53: "Murray, J. H. (1997). Hamlet on the Holodeck: The Future of Narrative in Cyberspace. MIT Press.",
    57: "Audio from Freesound: https://freesound.org/s/753178/ ; https://freesound.org/s/249156/ ; https://freesound.org/s/473525/ ; https://freesound.org/s/643766/ ; https://freesound.org/s/216923/ ; https://freesound.org/s/390549/ ; https://freesound.org/s/377206/",
    58: "Clown character: Mixamo, https://www.mixamo.com/#/?page=2&type=Character",
    59: "Picking up animation: Mixamo, https://www.mixamo.com/#/?page=1&type=Motion%2CMotionPack",
    60: "Unity Asset Store: Stylized Cracking Breakable Glass, https://assetstore.unity.com/packages/vfx/stylized-cracking-breakable-glass-366374",
    61: "Unity Asset Store: Free Stylized Sci-Fi Scout Bot, https://assetstore.unity.com/packages/3d/characters/robots/free-stylized-sci-fi-scout-bot-371626",
    62: "[TO VERIFY AND CITE: medal, clock, cell phone, BrickToy assets, all memory images, and any new assistant, memory, ambience, heartbeat, rumble, footsteps, monster, or swallow audio files added after A1. Remove any citation for an asset that is not used in the submitted build.]",
    64: (
        "Google AI Studio Nano Banana 2 Lite was used for image generation. Prompts requested "
        "horizontal Disney/Pixar-style 3D cartoon images of a tired young office worker at night "
        "with a thermos and medicine, then the same person sleeping in bed at night while "
        "maintaining character consistency."
    ),
}


CHINESE_REVISIONS = {
    5: "《记忆整理师》是一段关于修复焦虑梦境者潜意识的短篇沉浸式互动故事。玩家扮演一名刚入职的记忆整理师，执行自己的第一项任务。在办公室里，机器人助手介绍工作，并通过水晶球展示正在睡觉的梦境者。进入梦境者的潜意识后，玩家在一个以海马体为灵感的记忆空间中整理近期记忆。在助手的引导下，玩家查看三段近期记忆，并把它们放入“焦点”“语境”或“背景”注意区域。任务看似受到控制时，梦境世界突然破裂，一个由梦境者的焦虑和内化压力形成的巨型小丑出现。玩家被吞下，进入巨型小丑内部由痛苦记忆形成的空间。四场镜子对话帮助梦境者面对与这些记忆相连的信念。返回记忆空间后，玩家可以把全部七段记忆重新安排到三个注意区域。最终报告会反映这些选择，任务以玩家返回办公室结束。",
    13: "图 1 展示了故事如何在允许玩家选择的同时遵循 Freytag 模型。办公室选项会改变玩家开始任务的节奏，但都会回到同一个主要任务。三段近期记忆会引出巨人危机。在巨人内部，四面镜子可以按任意顺序访问；无帮助的回答会让该镜子的对话回到起点。四段记忆全部释放后，玩家返回记忆空间并重新安排七段记忆。因此，故事保持清晰的开端、危机、高潮、下降行动和结局，同时在每个阶段内部提供选择。",
    16: "办公室为玩家提供了规模较小但清晰的能动性。玩家可以询问工作内容、探索房间、返回助手身边，并决定何时开始。这些选项会形成不同的第一天节奏，但不会改变来访者或移除主要任务。这样既能控制开场的制作规模，也能让玩家感觉自己是一名正在选择如何准备的新员工。",
    17: "最强的分支出现在巨型小丑内部。四面镜子可以按任意顺序访问，仅顺序就有 24 种可能。每面镜子包含两轮对话。建设性回答会推动对话继续，有害的捷径则会强化负面声音，并让镜子回到起始状态。玩家可以再次尝试，因此失败表现为梦境者持续存在的自我怀疑，而不是传统的游戏结束画面。",
    18: "最终记忆安排为玩家提供了最广泛的选择。七段记忆都可以放入“焦点”“语境”或“背景”，玩家可以在确认前修改布局。这会产生 3^7，也就是 2,187 种注意模式。报告会根据每段记忆的最终区域改变。“焦点”不一定是好结果，“背景”也不会删除记忆；玩家决定的是梦境者现在应该关注什么。",
    19: "有些选择会被有意限制。第三段近期记忆放置后，巨人危机一定会开始，因为故事必须让痛苦记忆被面对，才能到达高潮。玩家不能避开这个事件，但可以选择查看顺序、修改位置、重试镜子对话，并产生真正不同的最终报告。因此，这段体验把固定的戏剧结构与能够改变到达方式和结局解释的选择结合起来。",
    20: "在编排好的抓取段落中，玩家的移动会暂时受到限制，但用户仍然可以在 3D 空间中自由观察。这会减少能动性，但符合叙事和物理逻辑：如果玩家已经被巨人抓住并吞下，就不应该还能自由走开。因此，能动性的降低成为恐惧、压力和失去控制这一具身体验的一部分。",
    22: "故事的主要规则是：记忆整理师可以移动注意力，但不能删除、改写或创造记忆。七个物件都与梦境者的一种恐惧相连：个人价值取决于完美工作、承受压力和不给别人添麻烦。水瓶、夕阳照片、积木、奖牌、闹钟、旧手机和红色批改笔展示了这一信念的不同侧面。因为它们属于同一个情绪问题，即使玩家以不同顺序查看，故事仍然保持连贯。",
    23: "故事结构遵循 Freytag 模型，同时针对互动体验进行了调整。玩家先了解工作，然后完成一项可控制的记忆任务，接着在小丑打断时失去控制。四场镜子对话形成高潮，因为它们移除了支撑巨人的信念。最终记忆安排属于下降行动，玩家在这里应用自己获得的理解。办公室报告通过展示玩家选择的结果完成结局。",
    24: "多个门控机制让顺序保持清晰。只有玩家完成办公室简报并阅读来访者报告后，水晶球才会出现。只有三段不同的近期记忆都被放置后，危机才会开始。只有第二轮获得建设性回答后，镜子才算完成；只有七段记忆都有区域时，最终报告才能确认。这些规则防止故事状态出错，同时保留每个阶段内部的自由。",
    25: "同一套结构也让项目保持可制作。四面镜子共用一个对话控制器，前期和最终放置阶段复用同一个记忆放置系统。最终报告为七个物件各写三个结果，因此 21 个短文本片段就能支持 2,187 种组合。这样可以提供大量结局，而不需要为每一种安排分别写一篇完整故事。",
    27: "办公室被设计成温暖的来访者关怀空间，而不是临床实验室。Ryan 对空间与地方的区分，使这里不只是一个中性容器，而是具有情绪和叙事意义的环境（Ryan, n.d.）。中央桌子为玩家提供清晰的方向点，沙发、娱乐物件、装饰品和生活细节则让房间不那么临床化。这些物件支持了这里用于关怀、交流和温和记忆修复的故事设定。",
    28: "海马体记忆空间被设计成一个由来访者强烈记住的物件组合而成的房间。因此，空间中的物品不会严格整齐排列，而更像被保留在记忆中的日常碎片。中央桌子引导玩家的身体和注意力：玩家走向桌子，拿起记忆物件，聆听其中的内容，然后把它放入“焦点”“语境”或“背景”。在最终阶段，同一个房间允许玩家比较并重新安排全部七个物件。",
    29: "小丑出现时，环境扩展成带有红雾、暗色天空和危险尺度的敌对外部空间。巨手撕开屋顶让玩家感到渺小和无力，支持了焦虑压倒有序记忆系统的故事时刻。在小丑内部，石墙、火把、锁链、骨头和四个镜子结构创造出被困住的沉重气氛。镜子可以按任意顺序接近，因此这里是一个枢纽，而不是单向走廊。镜子全部解决后，玩家返回受损的记忆空间；破损屋顶继续保留为危机留下的证据。",
    31: "感官设计为不同叙事阶段使用三个独立的 Global Volume，而机器人助手的声音连接整段体验。助手解释任务，在转场中减少困惑，并在视觉变得不稳定时让玩家继续理解剧情。",
    32: "在办公室到记忆空间的转场中，玩家把手放在水晶球上。这会触发持续的触觉振动，随后画面逐渐模糊并淡出为白色。温暖的办公室气氛转变成一个感官门槛，使进入来访者记忆空间的过程显得有意且可信。在 WebGL 版本中，点击会启动相同转场，但没有手柄触觉。",
    33: "在记忆阶段，独立的 Global Volume 为潜意识提供蓝紫色色调、漂浮粒子和更强的泛光。粒子让空气具有梦境感，聚光灯则强调记忆物件和注意区域。拿起物件还会播放对应的记忆叙述，因此图像、声音和物理操作会把注意力引向同一个故事片段。",
    34: "在危机阶段，噩梦 Global Volume 会降低曝光并加入红色暗角。巨人的脚步声、强烈心跳、耳鸣、屋顶被掀开的声音和暗色环境音共同加强紧张感。巨型小丑吞下玩家后，画面变黑，但吞咽声和助手的声音会继续，使玩家仍然理解当前情况。在巨人内部，每面未解决的镜子都会重复负面陈述；随着玩家解决镜子，房间会逐渐安静。",
    36: "水晶球：这个具有幻想感的物件让进入潜意识的转场显得自然。它没有采用纯技术界面，而是暗示观察、诊断和进入隐藏心理内容。记忆物件：桌面上的实体物件把不可见的心理事件变成玩家可以接近和拿取的东西，使记忆整理通过身体操作被理解，而不只依赖对话说明。",
    37: "记忆画布：玩家拿起物件时，记忆画布会显示具体的记忆片段。图像和叙述把日常物件与梦境者的经历连接起来。巨型小丑：小丑代表梦境者的负面情绪、内化压力和对评价的恐惧。夸张尺度把焦虑表现成已经变得过大、并开始打断修复工作的东西。镜子重复梦境者的有害信念，只有这些信念被重新理解后才会破裂。",
    38: "玩家被吞下后，助手的声音、吞咽音效和穿过管道坠落的图像共同传达一个清晰的新情况：我已经被吞下，现在位于巨人的身体内部。“焦点”“语境”和“背景”也作为符号发挥作用。它们描述记忆与当前注意力的距离，而不会把记忆称为好或坏。痛苦记忆可以移到背景而不被删除，积极记忆也可以移到焦点而不成为唯一正确答案。",
    40: "具身动作主要围绕两个行动者：玩家和巨型小丑。玩家的关键动作包括触碰水晶球、拿取记忆物件、把物件放入注意区域、被巨型小丑抓住、接近镜子、选择对话回应，以及重新安排全部七段记忆。这些动作使用具身隐喻，让玩家通过身体理解抽象过程（Antle et al., 2009）。注意力变成距离和位置，面对有害信念则变成接近镜子并与其对话。",
    41: "在编排好的抓取段落中，玩家的移动会暂时受到限制，但用户仍然可以在 3D 空间中自由观察。这会减少能动性，但符合叙事和物理逻辑：如果玩家已经被巨人抓住并吞下，就不应该还能自由走开。因此，能动性的降低成为恐惧、压力和失去控制这一具身体验的一部分。",
    42: "巨型小丑的动作构成危机。它撕开屋顶、伸入安全的工作空间、抓住玩家并将其吞下。这些动作推动剧情，同时制造强烈的视觉和听觉冲击。机器人助手也在故事内部行动：它向玩家简报、对危机作出反应，在四面镜子全部解决后恢复返回信号，并把玩家和释放出的物件带回记忆空间。",
    62: "[需要核对并引用：奖牌、闹钟、手机、BrickToy 资产、全部记忆图像，以及 A1 之后新增的助手、记忆、环境、心跳、隆隆声、脚步、怪物或吞咽音频。删除最终构建中未使用资产的引用。]",
    64: "Google AI Studio Nano Banana 2 Lite 被用于图像生成。提示词要求生成横向的 Disney/Pixar 风格 3D 卡通图像：一名疲惫的年轻办公室职员在夜间与保温杯和药物一起工作，随后是同一个人在夜间床上睡觉，并保持角色一致性。",
}


EXTRA_REFERENCES = [
    "Ryan, M.-L. (n.d.). Space, place and story. Course reading PDF.",
    "Slater, M. (2009). Place illusion and plausibility can lead to realistic behaviour in immersive virtual environments. Philosophical Transactions of the Royal Society B, 364(1535), 3549-3557. https://doi.org/10.1098/rstb.2009.0138",
]


EXTRA_ASSETS = [
    "Unity Asset Store: Lite Dungeon Pack Low Poly 3D Art by Gridness, https://assetstore.unity.com/packages/3d/environments/dungeons/lite-dungeon-pack-low-poly-3d-art-by-gridness-242692",
    "Unity Asset Store: Raft on the Desert Free Low Poly Pack, https://assetstore.unity.com/packages/3d/environments/landscapes/raft-on-the-desert-free-low-poly-pack-141948",
    "Unity Asset Store: Low Poly Dungeon Lite Fantasy Modular Kit, https://assetstore.unity.com/packages/3d/environments/dungeons/low-poly-dungeon-lite-fantasy-modular-kit-224313",
    "Unity Asset Store: Skybox Series Free, https://assetstore.unity.com/packages/2d/textures-materials/sky/skybox-series-free-103633",
    "EZPZ Interaction Toolkit by Matt Cabanag: https://github.com/mattcabanag ; Unity XR Interaction Toolkit and standard Unity packages were used through the project manifest.",
]


EXTRA_AI = [
    (
        "Deevid.ai was used for video generation. The prompt requested a Pixar-inspired single "
        "continuous shot of the young office worker at a night desk, noticing medicine, swallowing "
        "pills, drinking warm water, and returning to the computer, with consistent character "
        "design, simple office background, static camera, natural movement, and cinematic lighting.",
        "Deevid.ai 被用于视频生成。提示词要求制作一个受 Pixar 启发的连续镜头：年轻办公室职员坐在夜间办公桌前，注意到药物、吞下药片、喝温水并返回电脑；角色设计保持一致，办公室背景简单，相机固定，动作自然，并使用电影化灯光。",
    ),
    (
        "ElevenLabs was used to generate the robot assistant voice audio from the written dialogue "
        "script, so the assistant narration could remain clear and consistent across scene transitions.",
        "ElevenLabs 被用于根据书面对话脚本生成机器人助手的语音，使助手叙述在不同场景转场中保持清晰和一致。",
    ),
    (
        "OpenAI Codex was used for code assistance and report support. Prompts asked for "
        "implementation help based on current scripts and process analysis, debugging suggestions, "
        "rubric comparison, structure revision, and concise wording. Final design decisions, editing, "
        "asset selection, and submission responsibility remain my own.",
        "OpenAI Codex 被用于代码辅助和报告支持。提示词要求它根据当前脚本和流程分析提供实现帮助、调试建议、评分标准对照、结构修改和精简措辞。最终设计决定、编辑、资产选择和提交责任仍由我本人承担。",
    ),
]


FIGURES = [
    (27, ["image1.png"], [10.6], "Figure 3. Prototype view of the warm office care space. / 图 3：温暖办公室关怀空间的原型画面。", ["Warm office care space"]),
    (28, ["image2.png"], [10.6], "Figure 4. Prototype view of the hippocampus-inspired memory room and central table. / 图 4：海马体记忆空间及中央桌子的原型画面。", ["Hippocampus-inspired memory room"]),
    (29, ["image3.png", "image4.jpeg"], [7.1, 7.1], "Figure 5. Prototype views of the Giant's exterior and the enclosed memory space inside its body. / 图 5：巨人外部及其体内封闭记忆空间的原型画面。", ["Giant exterior", "Memory space inside the Giant"]),
    (32, ["image5.jpeg"], [10.6], "Figure 6. Touching the crystal ball begins the office-to-memory transition. / 图 6：触碰水晶球开始从办公室进入记忆空间的转场。", ["Crystal ball transition"]),
    (33, ["image6.jpeg"], [11.0], "Figure 7. Lighting, bloom, and particles distinguish the subconscious memory stage. / 图 7：灯光、泛光和粒子区分潜意识记忆阶段。", ["Subconscious memory stage"]),
    (34, ["image7.jpeg"], [11.0], "Figure 8. The nightmare volume and Giant scale make the crisis physically threatening. / 图 8：噩梦视觉和巨人尺度让危机具有身体威胁感。", ["Giant crisis in nightmare lighting"]),
    (38, ["image8.png"], [10.6], "Figure 9. The falling image clarifies the transition into the Giant's body. / 图 9：坠落图像说明玩家正在进入巨人身体内部。", ["Falling transition into the Giant"]),
    (42, ["image9.jpeg"], [10.6], "Figure 10. The Giant's close-up turns anxiety and loss of control into an embodied climax. / 图 10：巨人特写把焦虑和失去控制转化为具身高潮。", ["Giant close-up during capture"]),
]


def set_paragraph_text(paragraph, text):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run(text)


def extract_a1_media():
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    wanted = {name for _, names, _, _, _ in FIGURES for name in names}
    with ZipFile(A1_REPORT) as package:
        for name in wanted:
            target = MEDIA_DIR / name
            target.write_bytes(package.read(f"word/media/{name}"))


def add_paragraph_before(reference, text, chinese=None):
    paragraph = reference.insert_paragraph_before(text, style="Normal")
    if chinese:
        base.append_chinese(paragraph, chinese)
    return paragraph


def insert_paragraph_after(reference, text, chinese=None):
    paragraph = reference._parent.add_paragraph(text, style="Normal")
    reference._p.addnext(paragraph._p)
    if chinese:
        base.append_chinese(paragraph, chinese)
    return paragraph


def insert_figure_after(reference, names, widths, caption_text, alt_titles):
    image_paragraph = reference._parent.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True

    for index, (name, width, alt_title) in enumerate(zip(names, widths, alt_titles)):
        if index:
            image_paragraph.add_run("  ")
        shape = image_paragraph.add_run().add_picture(str(MEDIA_DIR / name), width=Cm(width))
        base.add_alt_text(shape, alt_title, alt_title)

    caption = reference._parent.add_paragraph(caption_text, style="Figure Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        base.set_east_asia(run)

    reference._p.addnext(image_paragraph._p)
    image_paragraph._p.addnext(caption._p)


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def remove_table(table):
    parent = table._element.getparent()
    parent.remove(table._element)


def update_summary_table(table):
    values = {
        (1, 1): "https://zoeyyy77.itch.io/sleeporganizer",
        (2, 1): "https://github.com/ZoeyW-zq/9903_Zoey",
        (3, 1): "scene_WebGL",
        (4, 1): "Zoey Wu z5658626",
    }
    for (row_index, cell_index), value in values.items():
        set_paragraph_text(table.rows[row_index].cells[cell_index].paragraphs[0], value)


def main():
    extract_a1_media()
    doc = Document(base.SOURCE)

    original_paragraphs = list(doc.paragraphs)
    status_paragraphs = [original_paragraphs[index] for index in (6, 7, 8, 9)]
    status_table = doc.tables[2]
    ai_table = doc.tables[5]
    reference_placeholder = original_paragraphs[54]
    asset_placeholder = original_paragraphs[62]
    ai_first = original_paragraphs[64]
    figure_targets = {index: original_paragraphs[index] for index, *_ in FIGURES}
    ending_target = original_paragraphs[21]

    for index, text_value in ENGLISH_REVISIONS.items():
        set_paragraph_text(original_paragraphs[index], text_value)

    for index, chinese in base.HEADINGS.items():
        if index != 6:
            base.make_bilingual_heading(original_paragraphs[index], chinese)

    translations = dict(base.PARAGRAPHS)
    translations.update(CHINESE_REVISIONS)
    for index, chinese in translations.items():
        if index in (7, 9):
            continue
        base.append_chinese(
            original_paragraphs[index],
            chinese,
            placeholder=index in {43, 45, 54, 62},
        )

    update_summary_table(doc.tables[1])
    table_translations = {key: value for key, value in base.TABLES.items() if key in (0, 1, 3, 4)}
    table_translations[1] = [
        ["提交字段", "草稿值"],
        ["WebGL 链接", "可公开游玩的 WebGL 链接"],
        ["GitHub 仓库", "GitHub 仓库链接"],
        ["首个场景", "scene_WebGL"],
        ["学生姓名 / zID", "Zoey Wu z5658626"],
    ]
    for table_index, translations_rows in table_translations.items():
        table = doc.tables[table_index]
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if row_index >= len(translations_rows) or cell_index >= len(translations_rows[row_index]):
                    continue
                value = translations_rows[row_index][cell_index]
                paragraph = cell.paragraphs[-1]
                placeholder = "待补充" in value or "插入英文截图" in value
                base.append_chinese(paragraph, value, placeholder=placeholder, size=8.2)

    for reference in EXTRA_REFERENCES:
        add_paragraph_before(reference_placeholder, reference)

    for asset in EXTRA_ASSETS:
        add_paragraph_before(asset_placeholder, asset)

    cursor = ai_first
    for english, chinese in EXTRA_AI:
        cursor = insert_paragraph_after(cursor, english, chinese)
        cursor.paragraph_format.keep_together = True

    base.replace_flowchart(doc)
    flowchart_heading = original_paragraphs[10]
    flowchart_heading.paragraph_format.page_break_before = True
    flowchart_heading.paragraph_format.keep_with_next = True
    flowchart_paragraph = next(p for p in doc.paragraphs if p._p.xpath(".//w:drawing"))
    flowchart_paragraph.paragraph_format.page_break_before = False
    for index in (16, 20, 31, 42):
        original_paragraphs[index].paragraph_format.keep_together = True
    base.insert_endings_figure(doc, ending_target)

    for index, names, widths, caption_text, alt_titles in FIGURES:
        insert_figure_after(figure_targets[index], names, widths, caption_text, alt_titles)

    for paragraph in status_paragraphs:
        remove_paragraph(paragraph)
    remove_table(status_table)
    remove_table(ai_table)

    base.update_footer(doc)
    doc.save(base.OUTPUT)
    print(base.OUTPUT)


if __name__ == "__main__":
    main()
