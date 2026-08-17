from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(r"E:\GitHub\9903_Zoey")
SOURCE = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer_Bilingual_Actual_Outcomes.docx"
OUTPUT = ROOT / "output" / "docx" / "DDES9903_A2_Report_Draft_Memory_Organizer_Bilingual_Actual_Outcomes_Stage3_Revised.docx"

CN_FONT = "Microsoft YaHei"
CN_GREY = RGBColor(0x58, 0x62, 0x70)


REVISIONS = {
    5: (
        "Memory Organizer is a short immersive interactive story about repairing an anxious Dreamer's relationship with memory and attention. The player begins as a newly hired Memory Organizer in a warm office, enters a hippocampus-inspired memory room, and provisionally places three recent memories in Focus, Context, or Background. This apparently manageable task is interrupted when a giant clown tears open the room, captures the player, and swallows them. The second half then changes both the setting and the rules: after a black-screen swallow and controlled fall, the player regains movement inside an enclosed Mirror Chamber built from stone, firelight, chains, bones, and four speaking mirrors. The mirrors hold a second-place medal, old alarm clock, old phone, and red correction pen, each connected to a different painful belief about achievement, rest, help-seeking, or mistakes. The player can approach the mirrors in any order and must navigate two-round conversations in which constructive responses release the memory, while harmful shortcuts strengthen the internal voice and restart that encounter. When all four echoes have been quieted, the body-space breaks down and the player returns with the released objects to the memory room. The final redistribution of all seven memories produces a compositional report, followed by a neutral Close Session ending rather than a single scored success state.",
        "《记忆整理师》是一段关于修复焦虑梦境者与记忆及注意力关系的短篇沉浸式互动故事。玩家以新入职的记忆整理师身份从温暖的办公室出发，进入以海马体为灵感的记忆空间，并把三段近期记忆暂时放入“焦点”“语境”或“背景”。这个看似可控的任务被巨型小丑打断：它撕开房间、抓住玩家并将其吞下。故事后半段同时改变了场景与规则：经过黑屏吞咽和受控坠落后，玩家在一个由石墙、火光、锁链、骨头和四面会说话的镜子组成的封闭镜室中重新获得移动能力。四面镜子分别封存第二名奖牌、旧闹钟、旧手机和红色批改笔，对应成就、休息、求助和犯错所引发的四种痛苦信念。玩家可以按任意顺序接近镜子，并完成两轮对话；建设性回应会释放记忆，有害的捷径则会强化内部声音并重启该次对话。当四段回声全部安静后，身体内部空间开始解体，玩家带着释放出的物件回到记忆空间。最后，全部七段记忆的重新分配会生成组合式报告，并以中性的“Close Session”结尾，而不是给出单一的成功分数。",
    ),
    9: (
        "Figure 1 maps the implemented state sequence onto Freytag's dramatic arc. OfficeDialogue and the initial memory placement establish the role and its attention-based rule, while GiantCrisis and SwallowTransition form a deliberately non-branching rupture. During this capture, control is reduced to looking, so the player's loss of navigation is part of the conflict rather than a menu choice. The black transition then moves the player to the pipe and controlled fall before GameStateController enters MirrorChamber and restores locomotion. The chamber is the interactive climax: four mirrors can be visited in any of 24 orders, and each contains a local dialogue loop. A harmful response in either round plays the negative reply, closes the interface, restores the opening voice, and leaves the mirror available for another attempt; it does not advance the global state or create a game-over. A constructive second-round response automatically shatters that mirror, permanently stops its echo, and releases the original object. Only the fourth completed mirror triggers the chamber's collapse and the transfer back to the memory room. FinalMemoryPlacement and the generated office report then provide falling action and denouement, making the second half a sustained playable stage rather than a short transition between the initial arrangement and ending.",
        "图 1 把当前实现的状态顺序映射到 Freytag 戏剧弧线。OfficeDialogue 与初次记忆放置建立角色及“以注意力整理记忆”的规则，而 GiantCrisis 与 SwallowTransition 构成一个有意不分支的断裂。在抓取过程中，玩家只保留观察能力，因此导航权的丧失属于冲突本身，而不是菜单选择。随后，黑屏转场把玩家送到管道与受控坠落段落；GameStateController 进入 MirrorChamber 状态后再恢复移动。镜室构成可游玩的高潮：四面镜子共有 24 种访问顺序，每面镜子内部都有一个局部对话循环。任一轮选择有害回应都会播放负面回复、关闭界面、恢复开场声音，并让镜子可以重新尝试；它不会推进全局状态，也不会触发 Game Over。第二轮的建设性回应会让镜子自动破裂、永久停止其回声，并释放原始物件。只有第四面镜子完成时，空间才会解体并把玩家送回记忆空间。FinalMemoryPlacement 与办公室中的生成式报告随后承担下降行动与结局，因此故事后半段是一个持续可玩的阶段，而不是初次分配与结局之间的短暂转场。",
    ),
    12: (
        "Agency is deliberately redistributed across the story rather than held at one constant level. The office offers low-risk preparation: the player can ask about the job, explore, return to the assistant, and decide when to begin. The initial memory room then introduces embodied classification through three recent objects. These early choices establish the player's role, but they are provisional; after the third distinct object is placed, the Giant crisis begins regardless of the chosen zones. This convergence is necessary because the later body-space addresses memories that the initial tidy surface has excluded.",
        "玩家能动性并不是在整个故事中保持同一强度，而是被有意重新分配。办公室提供低风险的准备空间：玩家可以询问工作内容、探索、返回助手身边，并决定何时开始。初始记忆空间则通过三个近期物件引入具身分类。这些前期选择用来建立玩家的角色，但都只是临时安排；第三个不同物件完成放置后，无论选择了哪个区域，巨人危机都会开始。这样的汇流是必要的，因为后续身体内部空间要处理的是被初始整洁表层排除在外的记忆。",
    ),
    13: (
        "The strongest procedural agency appears after the player lands inside the Giant. The Mirror Chamber is a hub rather than a corridor, so the medal, clock, phone, and red-pen mirrors can be approached in any order, producing at least 4! = 24 spatial sequences. Entering a mirror's proximity trigger reveals Start Conversation; beginning it opens a two-round exchange instead of immediately awarding progress. Round 1 presents three choices: two constructive interpretations lead to different Round 2 prompts, while one harmful shortcut fails immediately. Each Round 2 branch offers one response that completes the memory and one that repeats the damaging belief. This design makes the dialogue content playable: the player must distinguish growth from perfectionism, rest from irresponsibility, legitimate help-seeking from weakness, and careful correction from total incompetence.",
        "最强的程序性能动性出现在玩家落入巨人体内之后。镜室是枢纽而不是走廊，因此奖牌、闹钟、手机和红笔镜子可以按任意顺序接近，仅空间顺序就至少产生 4! = 24 种路径。进入某面镜子的接近触发区后，界面才显示“Start Conversation”；开始交谈会进入两轮对话，而不是立刻给予进度。第一轮提供三个选项：两个建设性解释分别进入不同的第二轮问题，一个有害捷径则立即失败。每条第二轮分支又包含一个完成记忆的回应和一个重复伤害性信念的回应。由此，对话内容本身成为玩法：玩家必须区分成长与完美主义、休息与不负责任、合理求助与软弱，以及谨慎修正与“完全无能”的自我判断。",
    ),
    14: (
        "Failure in the chamber is recoverable but perceptible. It plays the mirror's negative response, reinforces the local voice and atmosphere, closes the dialogue, and restores the Start Conversation state. The player remains physically present in the chamber and may retry immediately or walk to another unresolved mirror. Success is also expressed through the world rather than a score: the glass shatters automatically, the spatial echo stops, and the original memory object is released. The player cannot bypass the conversation by manually breaking the mirror. This creates consequential local choices without allowing one mistake to block the whole narrative or contaminate the final attention report.",
        "镜室中的失败可以恢复，但具有明确的感知后果。系统会播放该镜子的负面回应，强化局部声音与氛围，关闭对话，并恢复“Start Conversation”状态。玩家仍然身处镜室，可以立刻重试，也可以先走向另一面尚未解决的镜子。成功同样通过世界状态而不是分数表现：玻璃自动破裂，空间回声停止，原始记忆物件被释放。玩家不能通过手动打碎镜子绕过对话。这种设计让局部选择产生后果，同时避免一次错误阻断整条叙事，或错误地影响最终注意力报告。",
    ),
    15: (
        "The broad dramatic path remains fixed, but the experience does not pretend that every moment branches. The Giant will capture and swallow the player, all four deep memories must be released, and the chamber only collapses after the fourth completion. Within that fixed arc, the player controls mirror order, response path, retry timing, movement around the chamber, and the later arrangement of all seven memories. The final placement provides the widest expressive agency: seven independent Focus, Context, or Background assignments create 3^7 = 2,187 attention patterns, each assembled from the 21 one-sentence outcomes shown in Figure 2 rather than judged by a universal good/bad score.",
        "整体戏剧路径保持固定，但体验并不假装每个时刻都会分支。巨人必然抓住并吞下玩家，四段深层记忆都必须被释放，镜室也只有在第四次完成后才会解体。在这条固定弧线内部，玩家仍控制镜子顺序、回应路径、重试时机、在镜室中的移动，以及之后七段记忆的安排。最终放置提供最强的表达性能动性：七个彼此独立的“焦点”“语境”或“背景”选择产生 3^7，即 2,187 种注意模式；系统根据图 2 所示的 21 条单句结果组合报告，而不是用统一的好坏分数评判玩家。",
    ),
    16: (
        "Control changes are used as an agency rhythm. When the hand reaches for the player, translational movement is locked so the staged capture cannot be escaped, but head tracking or mouse look remains available and the camera is not forcibly rotated. The XR Origin follows the hand by position, the screen fades to black near the mouth, and SwallowController transfers the player to a controlled fall. Locomotion is restored at the chamber landing point, creating a clear shift from helpless observation back to exploration and decision-making. After all four mirrors resolve, control is briefly subordinated again to the collapse and transfer, then returns in the memory room for the final seven-object arrangement.",
        "控制权的变化形成了能动性的节奏。巨手伸向玩家时，系统会锁定位移，使已经编排的抓取无法被走位逃脱，但仍保留头部追踪或鼠标观察，相机也不会被强制旋转。XR Origin 只按位置跟随手部，接近嘴部时画面淡为黑色，SwallowController 随后把玩家转移到受控坠落段落。到达镜室落点后，移动恢复，体验由无力观察清晰转回探索与决策。四面镜子全部解决后，空间解体和转移会再次短暂接管控制；回到记忆空间后，玩家重新获得最终七物件安排的完整操作权。",
    ),
    22: (
        "Narrative coherence is protected by one ethical and causal rule: a Memory Organizer may redistribute attention but cannot delete, rewrite, or invent memories. The Giant's interior gives this rule a stronger test than the initial placement stage. Its four mirrors do not contain unrelated side stories; they externalise connected beliefs that maintain the same anxiety system. The medal repeats that only first place proves worth; the clock turns unfinished work into an alarm that makes rest feel irresponsible; the phone treats previous rejection as proof that help must never be requested; and the red pen converts one public mistake into evidence of permanent incompetence. These voices link the adult work crisis to older experiences, explaining why rearranging only three recent memories cannot resolve the problem.",
        "叙事连贯性由一条伦理与因果规则保护：记忆整理师可以重新分配注意力，但不能删除、改写或制造记忆。与初次放置相比，巨人体内对这条规则提出了更强的考验。四面镜子不是互不相关的支线故事，而是把维持同一焦虑系统的信念外化出来：奖牌重复“只有第一名才能证明价值”；闹钟把未完成工作变成警报，使休息显得不负责任；手机把曾经被拒绝解释为“以后绝不能求助”；红笔则把一次公开犯错转化为永久无能的证据。这些声音把成年后的工作危机连接到更早的经历，也解释了为什么只重新安排三段近期记忆无法解决问题。",
    ),
    23: (
        "The Freytag structure is sustained through a change in interaction grammar. Exposition uses conversation and observation; rising action uses object handling; the crisis removes locomotion and overwhelms the safe room; the climax restores exploration but replaces placement with dialogue navigation. Each successful mirror weakens the Giant by stopping one painful echo, so progress is audible and spatial before it is numerical. The assistant's reaction after the first release - that the mirror broke on its own and the echo disappeared - clarifies that reinterpretation, not violence, changes the space. After the fourth release, the assistant identifies that every painful echo is quiet, the body-space can no longer hold together, and transfer must begin. The final redistribution is therefore falling action: it applies the understanding gained in the chamber rather than introducing a second unrelated conflict.",
        "Freytag 结构通过交互语法的改变得以维持。开端使用交谈与观察，上升行动使用物件操作，危机剥夺移动并压倒安全房间，高潮则恢复探索，但把“放置物件”替换为“在对话中导航”。每成功解决一面镜子，就会停止一段痛苦回声，因此玩家首先从声音与空间变化中感知巨人被削弱，而不是从数字进度条中得知。第一段记忆释放后，助手指出镜子自行破裂且回声消失，从而明确是重新理解而不是暴力改变了空间。第四段释放后，助手说明所有痛苦回声已经安静、身体空间无法继续维持，并启动转移。因此，最终重新分配属于下降行动：它应用镜室中获得的理解，而不是引入第二个无关冲突。",
    ),
    24: (
        "State gates prevent the non-linear chamber from producing incoherent outcomes. The crisis begins only after three unique recent objects are placed. SwallowTransition changes the active scene root and moves the player before MirrorChamber restores movement. A mirror contributes to the global count only once and only after a constructive Round 2 response; a harmful response merely resets that mirror's local conversation. Completing one mirror cannot accidentally end the stage, and revisiting a failed mirror cannot duplicate progress. When the count reaches four, all unresolved loops have already stopped, the Nightmare ambience fades, the chamber collapses, and the four released objects are returned to explicit positions in the memory room with gravity and holdability restored. Only after all seven objects receive final zones can the report be confirmed.",
        "状态门控防止非线性镜室产生不连贯结果。危机只会在三个不同的近期物件都完成放置后开始。SwallowTransition 先切换激活的场景根并移动玩家，之后 MirrorChamber 才恢复移动。每面镜子只能贡献一次全局计数，而且必须在第二轮选择建设性回应；有害回应只重置该镜子的局部对话。完成一面镜子不会意外结束整个阶段，反复进入失败镜子也不会重复增加进度。当计数达到四时，所有未解决回声都已停止，Nightmare 氛围淡出，镜室解体，四个释放出的物件被送回记忆空间中明确的返回位置，并恢复重力与可持有状态。只有全部七个物件都获得最终区域后，报告才能确认。",
    ),
    25: (
        "The later stage also remains practical to implement because content and control are data-driven. Four differently authored mirror conversations share one MemoryDialogueController and the same proximity, choice, retry, completion, and auto-shatter logic. Their differences are configured as opening lines, audio clips, two Round 2 branches, and constructive or harmful outcomes rather than four bespoke code paths. The initial and final placement stages likewise share one phased controller. Finally, seven ternary placements select seven of 21 short report fragments, supporting 2,187 combinations without writing 2,187 complete endings. This reuse is not only technical economy: it keeps equivalent rules across all painful memories while allowing each belief to retain distinct dialogue and sound.",
        "后半段仍然具有可实现性，因为内容与控制逻辑采用数据驱动方式。四场内容不同的镜子对话共享同一个 MemoryDialogueController，以及相同的接近、选择、重试、完成与自动破裂逻辑。它们的差异通过开场台词、音频片段、两条第二轮分支，以及建设性或有害结果进行配置，而不是写成四套专用代码。初始与最终放置阶段也复用同一个分阶段控制器。最后，七个三元位置会从 21 个短报告片段中选择七个，从而支持 2,187 种组合，而不必写 2,187 个完整结局。这种复用不仅节省技术成本，也确保四段痛苦记忆遵循同一规则，同时保留各自不同的对话与声音。",
    ),
    27: (
        "The project uses three distinct scene roots - Office, Hippocampus, and Dungeon - so the second half is not spatially treated as an extension of the first room. The office is a warm, client-facing care space that frames memory work as a professional practice. The hippocampus room externalises attention through a central worktable and three physical placement zones. These familiar, readable layouts establish safety and orientation before the narrative breaks them. The crystal ball functions as a threshold between the administrative office and the Dreamer's subconscious, while the damaged memory-room roof later preserves evidence that the interior crisis has physically affected the previously ordered space.",
        "项目使用 Office、Hippocampus 和 Dungeon 三个独立场景根，因此后半段在空间上并不被当作前一个房间的延伸。办公室是温暖、面向来访者的关怀空间，把记忆工作框定为一种专业实践；海马体记忆空间则通过中央工作桌和三个实体放置区域，把注意力外化。两者熟悉且易读的布局先建立安全感与方向感，随后故事再将其打破。水晶球是行政办公室与梦境者潜意识之间的门槛，而后来受损的记忆空间屋顶则保留危机已经真实影响原有秩序的空间证据。",
    ),
    30: (
        "The initial memory room is organised around repeated pick-up, inspection, listening, and placement. Its open circulation lets the player compare recent memories from a stable central position, and the Focus, Context, and Background zones translate an abstract judgment into distance and location. After the body-space sequence, the same room returns with a changed meaning: four deep-memory objects now occupy explicit return points, can again respond to gravity and handling, and must be compared with the three recent objects. Reusing the room therefore creates narrative contrast rather than simple asset reuse; a space that first supported partial organisation becomes the site where the consequences of the Mirror Chamber are integrated.",
        "初始记忆空间围绕反复拿起、查看、聆听和放置来组织。开放的动线让玩家能够从稳定的中央位置比较近期记忆，“焦点”“语境”和“背景”区域则把抽象判断转化为距离与位置。身体内部段落结束后，同一个房间以不同意义回归：四个深层记忆物件出现在明确的返回点，再次响应重力并可以被拿取，还必须与三个近期物件一起比较。因此，重复使用该房间形成的是叙事对照，而不只是资产复用；最初只支持局部整理的空间，后来成为整合镜室后果的场所。",
    ),
    33: (
        "The Dungeon/Mirror Chamber uses a different spatial language and mechanic. Stone walls, torchlight, chains, bones, enclosed scale, and the bodily swallow transition replace the memory room's open work surface with a heavy trapped atmosphere. Four mirror structures are distributed as a hub, allowing the player to select direction and order after the forced fall. A conversation is not available from anywhere in the room: the player must physically approach a mirror's proximity zone before Start Conversation appears. Each mirror is therefore both a destination and a container for one internal voice. Because unresolved voices are spatial audio sources, orientation can be guided by sound as well as sight. On success, the mirror shatters automatically and releases its object into the chamber; after the fourth release, the space itself begins to fail. The environment thus performs narrative state changes instead of merely providing a backdrop for dialogue.",
        "Dungeon／镜室采用了不同的空间语言与游戏机制。石墙、火把、锁链、骨头、封闭尺度和具身化的吞咽转场，用沉重的受困氛围取代记忆空间开放的工作台。四个镜子结构以枢纽形式分布，使玩家在强制坠落之后重新选择方向与顺序。对话不能在房间任意位置启动；玩家必须身体接近某面镜子的触发区域，“Start Conversation”才会出现。因此，每面镜子既是一个空间目的地，也是一个内部声音的容器。由于未解决的声音使用空间音频源，玩家可以通过听觉和视觉共同判断方向。成功后，镜子会自动破裂，并把对应物件释放到镜室；第四段释放后，空间本身开始失效。环境由此直接展演叙事状态变化，而不只是对白的背景。",
    ),
    37: (
        "Sensory design distinguishes the three scene roots and then makes the Giant's interior progressively legible through sound. The office and hippocampus use their own Global Volume treatments, while the crisis cross-fades from the hippocampus volume into the Nightmare volume and begins a heartbeat. The robot assistant provides a continuous narrative voice across these visual discontinuities. Its warning - 'Wait - I heard something' - and escalating reaction precede the Giant's appearance; after the swallow, the line 'Well, that's so... intimate. And so loud' acknowledges both the bodily setting and the chamber's overlapping voices. The assistant therefore does more than explain controls: its timed responses translate changes in space, danger, and progress into character dialogue.",
        "感官设计先区分三个场景根，再通过声音让巨人体内的状态逐步变得可读。办公室与海马体各自使用不同的 Global Volume 处理，危机开始时则从海马体氛围淡出、淡入 Nightmare 氛围，并启动心跳声。机器人助手的声音跨越这些视觉断裂，保持叙事连续。它先以“Wait - I heard something”发出警告，再在巨人出现前逐步升级反应；被吞下之后，“Well, that's so... intimate. And so loud”同时确认了身体内部场景与多重声音重叠的状态。因此，助手不只是解释操作，其按时序出现的回应还把空间、危险与进度变化转译为角色对白。",
    ),
    38: (
        "The crisis audio follows a staged causal sequence rather than playing as one undifferentiated ambience. The assistant's warning is followed by footsteps; the system waits for those footsteps to finish before starting the roof rumble and Giant animation. The grab introduces a groan and a later post-groan, while the heartbeat and Nightmare processing maintain pressure. Translational movement is locked, but the player can still turn and observe the approaching hand and broken roof. Near the mouth, a black fade overlaps the final hand movement. Swallow sounds and the assistant's panicked 'OH NO NO NO' continue while vision is absent, so audio carries the event across the hidden teleport to the pipe and controlled fall. This sound-image sequencing makes a technical scene change readable as one continuous bodily journey.",
        "危机音频遵循分阶段的因果顺序，而不是作为一层无法区分的环境声同时播放。助手警告之后先出现脚步声；系统等待脚步结束，才启动屋顶震动与巨人动画。抓取时加入呻吟，之后还有第二段呻吟，而心跳与 Nightmare 视觉处理持续维持压力。位移被锁定，但玩家仍可以转头观察逼近的巨手与破裂屋顶。接近嘴部时，黑色淡出与最后一段手部运动重叠；视觉消失后，吞咽音效和助手惊慌的“OH NO NO NO”仍然继续，使声音跨越隐藏的管道传送与受控坠落。这样的声画时序把技术性的场景切换组织成一段连续的身体旅程。",
    ),
    41: (
        "Inside the Mirror Chamber, sound becomes both atmosphere and a progress system. Each unresolved mirror repeats its opening belief at approximately one-second intervals. The red pen says that public criticism proves a lack of ability; the phone insists that asking for help only leads to rejection; the alarm clock says unfinished work keeps chasing the Dreamer; and the medal claims that only first place would satisfy the parents. Dialogue sources are fully spatial, so four active voices produce an intentionally crowded internal soundscape. The opening loop continues when a conversation begins until the player commits to a response. A harmful response ends the attempt and restores the loop; a successful constructive response stops it permanently. The room therefore becomes audibly quieter one memory at a time.",
        "在镜室内部，声音同时承担氛围与进度系统。每面尚未解决的镜子都会以约一秒间隔重复其开场信念：红笔认为公开批评证明自己没有能力；手机坚持求助只会再次被拒绝；闹钟声称未完成工作会一直追赶梦境者；奖牌则认为只有第一名才能让父母满意。对话音源采用完全空间化设置，因此四段同时存在的声音会形成有意拥挤的内部声场。对话开始后，开场循环仍会继续，直到玩家真正选择回应。有害回应会结束本次尝试并恢复循环；建设性成功回应则永久停止该声音。因此，房间会随着每段记忆完成而逐步变得安静。",
    ),
    44: (
        "The chamber's echo is implemented differently for platform reliability. DialogueAudioSource uses full 3D spatial blending, while WebGLEchoAudio recreates an echo tail with up to two delayed AudioSources because AudioMixer DSP effects are unreliable in WebGL. The delayed copies begin at roughly 0.3-second intervals and decay in volume, making each belief seem to persist in the body-space. Success is marked by automatic glass breakage, the abrupt disappearance of that echo, and the assistant's surprised confirmation. After all four, the assistant announces that every painful echo is quiet and that the space is shifting; the Nightmare ambience fades as the chamber breaks down and the transfer begins. These cues still require final end-to-end WebGL and XR timing checks, so the report describes implemented sequencing rather than claiming completed platform validation.",
        "为了保证不同平台上的可靠性，镜室回声采用了不同的实现方式。DialogueAudioSource 使用完全 3D 空间混合；由于 AudioMixer 的 DSP 效果在 WebGL 中不可靠，WebGLEchoAudio 使用最多两个延迟 AudioSource 重建回声尾音。延迟副本以约 0.3 秒间隔开始并逐步衰减，使每个信念仿佛持续滞留在身体空间中。成功通过自动碎裂、该段回声突然消失，以及助手惊讶的确认共同标记。四段全部完成后，助手宣布所有痛苦回声已经安静、空间正在移动；随着镜室解体和转移开始，Nightmare 氛围逐渐淡出。这些线索仍需完成 WebGL 与 XR 的端到端时序检查，因此报告只描述已经实现的编排，不声称平台验证已经完成。",
    ),
    48: (
        "The crystal ball and the memory objects establish the project's first semiotic system: hidden mental content becomes something visible, audible, and physically handled. However, the later objects deepen this system because their ordinary forms carry specific internal rules. The second-place medal signifies the conversion of achievement into insufficient rank; the alarm clock turns time and unfinished work into a demand that never permits rest; the old phone materialises an unsent request for support; and the red correction pen magnifies one mistake into an enduring judgment of ability. Their release from the mirrors does not make the experiences disappear. The same objects return to the memory room and remain available for Focus, Context, or Background, preserving the distinction between changing attention and erasing history.",
        "水晶球与记忆物件建立了项目的第一套符号系统：隐藏的心理内容变成可见、可听、可被身体操作的东西。但后续物件进一步深化了这套系统，因为它们的日常形态承载具体的内部规则。第二名奖牌象征成就被转化为“不够高的名次”；闹钟把时间与未完成工作变成永远不允许休息的要求；旧手机把没有发出的求助实体化；红色批改笔则把一次错误放大为对能力的长期判决。它们从镜子中释放并不表示经历已经消失。同一批物件会回到记忆空间，继续可以被放入“焦点”“语境”或“背景”，从而保持“改变注意力”与“抹除历史”之间的区别。",
    ),
    49: (
        "The Giant's body is a semiotic environment rather than a neutral dungeon skin. Being swallowed literalises the feeling of being consumed by anxiety, while stone, chains, bones, firelight, and enclosed scale suggest that the Dreamer is trapped inside beliefs that have accumulated over time. The four mirrors transform internal speech into architecture: each apparently objective reflection repeats a harsh belief until the player answers it. The glass cannot be broken through attack. It shatters automatically after a constructive interpretation, indicating that the structure loses coherence when its premise is challenged. The released object remains intact, while the reflective prison disappears. This distinction shows that the memory is preserved even though its controlling interpretation has changed.",
        "巨人的身体是一个符号环境，而不是中性的地牢皮肤。被吞下把“被焦虑吞噬”的感受字面化；石头、锁链、骨头、火光与封闭尺度则暗示梦境者被长期累积的信念困住。四面镜子把内部语言转化为空间建筑：每个看似客观的反射会不断重复苛刻信念，直到玩家回应它。玻璃不能通过攻击打碎，而会在建设性解释之后自动破裂，表示当信念前提受到挑战时，该结构本身失去连贯性。释放出的物件保持完整，消失的是反射性的牢笼。这一区别说明记忆仍被保留，改变的是它对当下的控制性解释。",
    ),
    50: (
        "The swallow transition combines blackness, swallowing audio, assistant dialogue, and a falling pipe sequence into a readable sign of crossing into the body. Once inside, spatial sound makes the mirrors function like voices occupying different positions in the Dreamer's mind. Their gradual silence signifies progress without requiring a score display; the chamber's collapse after the fourth silence signifies that the Giant depended on these repeated beliefs for stability. Back in the memory room, Focus, Context, and Background continue the same non-moral semiotic logic. They represent attention distance, not good, neutral, and bad. A painful memory placed in Background may reduce immediate pressure but make its warning easier to overlook, while a positive memory in Focus may support recovery without becoming the only correct answer.",
        "吞咽转场把黑暗、吞咽音效、助手对白和管道坠落组合成一个清晰的“进入身体”符号。进入之后，空间声音让镜子像占据梦境者思维中不同位置的声音。它们逐步安静，不需要分数显示也能指示进度；第四段声音停止后镜室解体，则表明巨人的稳定依赖这些反复信念。回到记忆空间后，“焦点”“语境”和“背景”延续同一套非道德化符号逻辑：它们表示注意距离，而不是好、中、坏。痛苦记忆放入“背景”可能减轻即时压力，却也可能让警告更容易被忽视；积极记忆放入“焦点”可以支持恢复，但不会因此成为唯一正确答案。",
    ),
    54: (
        "Embodied action changes with each narrative space. In the office, the player explores and activates the crystal ball; in the first memory room, they pick up objects, listen while holding them, and place them into attention zones. Inside the Giant, object sorting is suspended. The player instead lands, reorients, follows overlapping spatial voices, walks into a mirror's proximity area, selects Start Conversation, and chooses responses across two rounds. The mirror does not accept a physical strike as a solution, so the player's embodied task is approach, attention, and speech rather than combat. After each successful conversation, an object is released but is only restored to normal holdable use when the story returns to the memory room. The final act recombines dialogue-derived understanding with physical sorting across all seven objects.",
        "具身动作会随叙事空间改变。在办公室，玩家进行探索并激活水晶球；在第一个记忆空间中，玩家拿起物件、持有时聆听内容，并把它们放入注意区域。进入巨人体内后，物件分类暂时停止。玩家需要落地、重新辨认方向、跟随重叠的空间声音、走入镜子的接近区域、选择“Start Conversation”，并在两轮对话中作出回应。镜子不接受物理攻击作为解决方式，因此玩家的具身任务是靠近、注意与对话，而不是战斗。每次成功对话后，一个物件会被释放，但只有故事返回记忆空间时才恢复正常的可持有状态。最后一个阶段把对话中获得的理解与七个物件的实体分类重新结合起来。",
    ),
    55: (
        "The capture and swallow create an enacted loss of control without switching to an external cinematic camera. The Giant tears away the roof, reaches into the room, aligns the player with the hand by moving the XR Origin, and carries them toward the animated mouth. Locomotion and jumping are disabled, but head direction or mouse look remains active, allowing the player to witness their changing scale and position. The black fade does not end the embodied sequence: swallowing audio continues, the player is moved to the pipe start point, and ControlledFall delivers them to the chamber landing point before locomotion is unlocked. The contrast between forced transport and restored movement makes the chamber feel like a new playable condition, not only a change of scenery.",
        "抓取与吞咽在不切换到外部电影镜头的情况下，展演了控制权的丧失。巨人掀开屋顶、把手伸入房间，通过移动 XR Origin 让玩家与手部对齐，再把玩家带向动画中的嘴部。位移与跳跃被禁用，但头部方向或鼠标观察仍然有效，使玩家能够目睹自身尺度与位置变化。黑色淡出并没有结束具身段落：吞咽音频继续，玩家被移动到管道起点，ControlledFall 把玩家送到镜室落点，之后才解除移动锁定。强制运输与恢复移动之间的对比，使镜室成为一种新的可游玩状态，而不只是换了布景。",
    ),
    56: (
        "The Giant, the assistant, and the environment all perform narrative actions. The Giant approaches, removes the roof, grabs, groans, and swallows; the assistant warns the player, reacts during the black transition, comments on the chamber's loud internal voices, recognises the first automatic shatter, counts the fourth release, and initiates the escape. The mirrors repeatedly speak until answered, then shatter and fall silent; after the last one, the body-space shifts and breaks down. On return, the assistant names the memory room as familiar solid ground, while the four freed objects reappear at deliberate positions. These enacted changes allow plot information to be delivered through bodies, voices, and space rather than through an abstract quest log.",
        "巨人、助手与环境都在执行叙事动作。巨人接近、移除屋顶、抓取、呻吟并吞咽；助手向玩家发出警告，在黑屏转场中作出反应，评论镜室内部声音的嘈杂，识别第一次自动破裂，在第四段释放时确认进度，并启动逃离。镜子在得到回应前持续发声，之后破裂并安静；最后一面完成后，身体空间移动并解体。返回时，助手把记忆空间称为熟悉而坚实的地面，四个释放出的物件则在预先设置的位置重新出现。这些展演变化使剧情信息通过身体、声音与空间传达，而不是依赖抽象任务日志。",
    ),
}


def set_east_asia(run, font_name=CN_FONT):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def replace_bilingual_paragraph(paragraph, english, chinese):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    paragraph.add_run(english)
    paragraph.add_run().add_break()
    cn = paragraph.add_run(chinese)
    set_east_asia(cn)
    cn.font.size = Pt(9.5)
    cn.font.color.rgb = CN_GREY
    cn.font.highlight_color = WD_COLOR_INDEX.AUTO


def replace_bilingual_cell(cell, english, chinese):
    paragraph = cell.paragraphs[0]
    replace_bilingual_paragraph(paragraph, english, chinese)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def update_state_table(doc):
    table = doc.tables[2]
    updates = {
        (4, 2): (
            "Movement is withheld during the staged capture, while looking remains available; black fade and controlled fall carry the transition.",
            "编排式抓取期间暂停位移但保留观察；黑屏与受控坠落完成空间转场。",
        ),
        (5, 2): (
            "Choose one of 24 mirror orders and navigate two-round dialogue; harmful responses reset locally, while success stops one spatial echo.",
            "选择 24 种镜子顺序之一并完成两轮对话；有害回应局部重置，成功则停止一段空间回声。",
        ),
        (6, 2): (
            "Reintegrate four released deep memories and revise seven independent Focus/Context/Background assignments.",
            "整合四段已释放的深层记忆，并修改七个独立的焦点／语境／背景分配。",
        ),
        (7, 1): (
            "BackToOffice / ViewReport / SessionComplete",
            "BackToOffice / ViewReport / SessionComplete",
        ),
        (7, 2): (
            "Read the compositional report, then use Close Session for the neutral spoken ending and final fade to black.",
            "阅读组合式报告，再以 Close Session 触发中性口述结尾并最终淡为黑屏。",
        ),
    }
    for (row_index, cell_index), (english, chinese) in updates.items():
        replace_bilingual_cell(table.rows[row_index].cells[cell_index], english, chinese)


def main():
    doc = Document(SOURCE)
    if len(doc.paragraphs) != 92 or len(doc.tables) != 4:
        raise RuntimeError(
            f"Unexpected source structure: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
        )

    for index, (english, chinese) in REVISIONS.items():
        replace_bilingual_paragraph(doc.paragraphs[index], english, chinese)
        doc.paragraphs[index].paragraph_format.keep_together = True

    update_state_table(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
